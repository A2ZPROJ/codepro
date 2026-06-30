# -*- coding: utf-8 -*-
import shutil, os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

SRC = r"\\2s-eng-servidor\maringa\PLANILHAS FINAIS\UBIRATÃ\PO-01\SBB-01\MD-048-GER-PB-PBEN-02-R0.docx"
WORK = r"C:\Users\lcabd\AppData\Local\Temp\claude\C--Users-lcabd\ea7ea8c4-9e05-4292-9d4b-335930a19f2c\scratchpad\PO01_estilo_md048.docx"
MAPA_LOC = r"C:\Users\lcabd\OneDrive - 2S ENGENHARIA DE AGRIMENSURA E GEOTECNOLOGIA\Área de Trabalho\_extracao\mapas\Mapa1_Localizacao.png"
MAPA_3D  = r"C:\Users\lcabd\OneDrive - 2S ENGENHARIA DE AGRIMENSURA E GEOTECNOLOGIA\Área de Trabalho\_extracao\mapas\Mapa6_3D_Topografia.png"
FALTA = "[A PREENCHER]"

shutil.copyfile(SRC, WORK)
doc = Document(WORK)

# ---------- run-aware replace (preserva formatacao do 1o run de cada match) ----------
def replace_in_paragraph(p, old, new):
    if old not in p.text:
        return 0
    runs = p.runs
    if not runs:
        return 0
    # mapeia (run_idx, offset) por posicao de caractere
    idx = []
    full = []
    for ri, r in enumerate(runs):
        for ci, ch in enumerate(r.text):
            idx.append((ri, ci)); full.append(ch)
    full = "".join(full)
    count = 0
    start = 0
    while True:
        pos = full.find(old, start)
        if pos < 0:
            break
        count += 1
        end = pos + len(old)
        sr, sc = idx[pos]
        er, ec = idx[end-1]
        # escreve new no run inicial, apaga o resto do trecho
        srun = runs[sr]
        if sr == er:
            t = srun.text
            srun.text = t[:sc] + new + t[ec+1:]
        else:
            srun.text = srun.text[:sc] + new
            for ri in range(sr+1, er):
                runs[ri].text = ""
            runs[er].text = runs[er].text[ec+1:]
        # recomputa (texto mudou)
        return count + replace_in_paragraph(p, old, new)
    return count

def all_paragraphs(document):
    for p in document.paragraphs:
        yield p
    for t in document.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    yield p
    for sec in document.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            try:
                for p in hf.paragraphs:
                    yield p
            except Exception:
                pass

# ---------- 1) IDENTIDADE PA1 -> PO1 (ordem: especifico -> generico) ----------
IDENT = [
    ("SB-PA1", "SB-PO1"),
    ("SB-B01", "SB-PO1-01"), ("SB-B02", "SB-PO1-02"), ("SB-B03", "SB-PO1-03"),
    ("SB-B04", "SB-PO1-04"), ("SB-B05", "SB-PO1-05"),
    ("Sub-Bacia PA1", "Sub-Bacia PO1"),
    ("Sub-bacia PA1", "Sub-bacia PO1"),
    ("sub-bacia PA1", "sub-bacia PO1"),
    ("bacia PA1", "bacia PO1"),
    ("PA1", "PO1"),
]
nident = 0
for p in all_paragraphs(doc):
    for old, new in IDENT:
        nident += replace_in_paragraph(p, old, new)
print("identidade PA1->PO1 substituicoes:", nident)

# ---------- helper: setar texto de uma celula com (ou sem) highlight ----------
def set_cell(cell, text, falta=False):
    # limpa
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    if falta:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return r

def cell_txt(cell):
    return cell.text.strip()

# ---------- localizar tabelas por assinatura ----------
def find_table(sig_first_cell, ncols=None):
    for t in doc.tables:
        try:
            if cell_txt(t.rows[0].cells[0]) == sig_first_cell and (ncols is None or len(t.columns) == ncols):
                return t
        except Exception:
            pass
    return None

tabs = doc.tables
print("total tabelas:", len(tabs))
for i, t in enumerate(tabs):
    print("  T%d %dx%d  r0c0=%r" % (i, len(t.rows), len(t.columns), cell_txt(t.rows[0].cells[0])[:30]))

# ---------- 2) T1 REVISOES -> nossa emissao ----------
T1 = tabs[1]
if cell_txt(T1.rows[0].cells[0]).upper().startswith("REVIS"):
    # cabecalho fica; primeira linha = nossa emissao; remove o resto
    rows = list(T1.rows)
    vals = ["R0", "EMISSÃO INICIAL", FALTA, "2S ENGENHARIA", FALTA]
    for ci, c in enumerate(rows[1].cells):
        v = vals[ci] if ci < len(vals) else ""
        set_cell(c, v, falta=(v == FALTA))
    for row in rows[2:]:
        row._element.getparent().remove(row._element)
    print("T1 revisoes ajustada")

# ---------- 3) T3 VAZOES (16 anos) -> esqueleto p/ preencher ----------
# mantem coluna Ano e Per capita(=138); zera o resto
T3 = None
for t in tabs:
    if "Pop. Total" in cell_txt(t.rows[0].cells[1]) or (len(t.columns) >= 12 and len(t.rows) >= 16):
        T3 = t; break
if T3 is not None:
    hdr_rows = 2
    # descobrir indice da coluna 'Per capita'
    pc_col = None
    for ci, c in enumerate(T3.rows[1].cells):
        if "capita" in cell_txt(c).lower() or "capita" in cell_txt(T3.rows[0].cells[ci]).lower():
            pc_col = ci
    for ri in range(hdr_rows, len(T3.rows)):
        for ci, c in enumerate(T3.rows[ri].cells):
            if ci == 0:
                continue  # mantem Ano
            if ci == pc_col:
                set_cell(c, "138")
                continue
            set_cell(c, "")  # vazio p/ preencher manual
    print("T3 vazoes -> esqueleto (Ano + per capita 138; resto vazio)")

# ---------- 4) T4 SOLEIRAS -> [A PREENCHER] ----------
T4 = None
for t in tabs:
    if "Soleiras" in cell_txt(t.rows[0].cells[1]) or "Micro Bacia" in cell_txt(t.rows[0].cells[0]):
        T4 = t; break
if T4 is not None:
    for ri in range(1, len(T4.rows)):
        for ci in range(1, len(T4.rows[ri].cells)):
            set_cell(T4.rows[ri].cells[ci], FALTA, falta=True)
    print("T4 soleiras -> [A PREENCHER]")

# ---------- 5) T5 PARAMETROS SewerGEMS: manter; so taxa linear -> [A PREENCHER] ----------
T5 = None
for t in tabs:
    if len(t.rows) < 2:
        continue
    if "Parâmetros" in cell_txt(t.rows[0].cells[0]) or "Dmín" in cell_txt(t.rows[1].cells[0]):
        T5 = t; break
if T5 is not None:
    for ri in range(1, len(T5.rows)):
        if "linear" in cell_txt(T5.rows[ri].cells[0]).lower() or "contribuição linear" in cell_txt(T5.rows[ri].cells[0]).lower():
            set_cell(T5.rows[ri].cells[1], FALTA, falta=True)
    print("T5 parametros -> taxa linear [A PREENCHER], resto mantido")

# ---------- 6) T6 REDE EXISTENTE -> [A PREENCHER] ----------
T6 = None
for t in tabs:
    if cell_txt(t.rows[0].cells[0]) == "Tipo de Rede":
        T6 = t; break
if T6 is not None:
    for ri in range(1, len(T6.rows)):
        for ci in range(len(T6.rows[ri].cells)):
            cur = cell_txt(T6.rows[ri].cells[ci])
            if cur and not cur.upper().startswith("TOTAL") and cur not in ("Rede Coletora",):
                # so troca os numeros/material por preencher onde houver extensao
                pass
    # mais simples: extensoes -> [A PREENCHER]
    for ri in range(1, len(T6.rows)):
        for ci in range(len(T6.rows[ri].cells)):
            cur = cell_txt(T6.rows[ri].cells[ci])
            if "3.775" in cur:
                set_cell(T6.rows[ri].cells[ci], FALTA, falta=True)
    print("T6 rede existente -> extensoes [A PREENCHER]")

# ---------- 7) T7 REDE PROJETADA POR DN -> PO-01 ----------
T7 = None
for t in tabs:
    r0 = [cell_txt(c) for c in t.rows[0].cells]
    if r0[:2] == ["Diâmetro", "Material"] and len(t.rows) >= 4:
        T7 = t; break
if T7 is not None:
    # r1 = DN150, r2 = DN200, r3 = (FD, remover), r4 = total
    set_cell(T7.rows[1].cells[0], "150 mm"); set_cell(T7.rows[1].cells[1], "PVC"); set_cell(T7.rows[1].cells[2], "34.916,39")
    set_cell(T7.rows[2].cells[0], "200 mm"); set_cell(T7.rows[2].cells[1], "PVC"); set_cell(T7.rows[2].cells[2], "299,60")
    # remover linha FD (penultima) se existir 5 linhas
    if len(T7.rows) == 5:
        rowFD = T7.rows[3]
        rowFD._element.getparent().remove(rowFD._element)
    # total (ultima)
    last = T7.rows[len(T7.rows)-1]
    set_cell(last.cells[len(last.cells)-1], "35.215,99")
    print("T7 rede projetada -> PO-01 (34.916,39 + 299,60 = 35.215,99)")

# ---------- 8) T8 DISPOSITIVOS PV/TL/TQ -> contagem real PO-01 ----------
# PV/TL = contagem manual por bacia (pranchas CAD): PV 685, TL 64.
# TQ = 35 (extraido das OSE, celula U12). TOTAL dispositivos = PV+TL (TQ nao soma).
# Ligacoes prediais (1.754) entram como linha informativa, fora do total.
import copy
T8 = None
for t in tabs:
    if "Dispositivos de Inspeção" in cell_txt(t.rows[0].cells[0]):
        T8 = t; break
if T8 is not None:
    tq_idx = total_idx = None
    for ri in range(1, len(T8.rows)):
        lbl = cell_txt(T8.rows[ri].cells[0]).upper()
        if lbl.startswith("PV"):
            set_cell(T8.rows[ri].cells[1], "685")
        elif lbl.startswith("TL"):
            set_cell(T8.rows[ri].cells[1], "64")
        elif lbl.startswith("TQ"):
            set_cell(T8.rows[ri].cells[1], "35"); tq_idx = ri
        elif "TOTAL" in lbl:
            set_cell(T8.rows[ri].cells[1], "749"); total_idx = ri
    # insere linha "Ligações Prediais" antes do TOTAL (copia o <tr> do TQ p/ herdar formato)
    if tq_idx is not None and total_idx is not None:
        new_tr = copy.deepcopy(T8.rows[tq_idx]._tr)
        T8.rows[total_idx]._tr.addprevious(new_tr)
        newrow = T8.rows[tq_idx + 1]
        set_cell(newrow.cells[0], "Ligações Prediais")
        set_cell(newrow.cells[1], "1.754")
    print("T8 dispositivos -> PV 685 / TL 64 / TQ 35 / Ligacoes 1.754 / TOTAL 749")

# ---------- 9) T9 PROFUNDIDADE -> faixas [A PREENCHER], total 35.215,99 ----------
T9 = None
for t in tabs:
    if "Profundidade" in cell_txt(t.rows[0].cells[0]):
        T9 = t; break
if T9 is not None:
    # extensao por faixa de profundidade, calculada das OSE (prof media por
    # segmento entre nos, col 19) e reconciliada ao total oficial 35.215,99
    FAIXAS_PROF = {
        "1,25": "5.014,31", "2,00": "23.872,43", "3,00": "4.862,75",
        "4,00": "1.242,76", "5,00": "223,74",
    }
    for ri in range(1, len(T9.rows)):
        lbl = cell_txt(T9.rows[ri].cells[0])
        up = lbl.upper()
        if up.startswith("EXTENSÃO TOTAL") or up.startswith("EXTENSAO TOTAL"):
            set_cell(T9.rows[ri].cells[1], "35.215,99")
            continue
        val = None
        for key, v in FAIXAS_PROF.items():
            if key in lbl:
                val = v; break
        set_cell(T9.rows[ri].cells[1], val if val else FALTA, falta=(val is None))
    print("T9 profundidade -> faixas das OSE (total 35.215,99)")

# ---------- 10) EDITS DE PROSA (numeros PA1-especificos) ----------
def para_replace_global(old, new, falta=False, maxn=99):
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            c = replace_in_paragraph(p, old, new)
            n += c
            if falta and c:
                for r in p.runs:
                    if new in r.text:
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            if n >= maxn:
                break
    return n

# extensao total na prosa
print("P174 ext:", para_replace_global("31.521,12 m", "35.215,99 m"))
# vazao maxima horaria 10,68 -> [A PREENCHER]
print("vazao 10,68:", para_replace_global("10,68 L/s", FALTA + " L/s", falta=True))
# indice de atendimento 100,00%
print("indice:", para_replace_global("é de: 100,00%", "é de: " + FALTA, falta=True))
# vazao pontual loteamento
para_replace_global("178 habitantes", FALTA + " habitantes", falta=True)
para_replace_global("vazão média de 0,33 L/s", "vazão média de " + FALTA + " L/s", falta=True)
para_replace_global("vazão máxima horária para final de plano de 0,51 L/s", "vazão máxima horária para final de plano de " + FALTA + " L/s", falta=True)
# travessia 05 / 46,23 / 1,48 / seis alternativas -> marca
para_replace_global("Travessia 05, que apresenta extensão aproximada de 46,23 m", "Travessia " + FALTA + ", que apresenta extensão aproximada de " + FALTA + " m", falta=True)
para_replace_global("profundidade de 1,48 m", "profundidade de " + FALTA + " m", falta=True)
# rede existente substituir
para_replace_global("02 trechos de rede coletora existente", FALTA + " trechos de rede coletora existente", falta=True)
para_replace_global("262,92 metros", FALTA + " metros", falta=True)
# referencia ao memorial de calculo MC-048
para_replace_global("MC-048-GER-PB-PBEN-02-R0", "Memorial de Cálculo do Projeto Executivo do SB-PO1 [A PREENCHER]", falta=True)

# legenda da Figura 3.1 (imagem agora e o modelo 3D, nao uma delimitacao 2D)
para_replace_global("Delimitação dos Subsistema de Contribuição denominado SB-PO1",
                    "Traçado da rede coletora do subsistema SB-PO1 sobre o modelo 3D do terreno")
# Fonte das figuras que trocamos -> 2S (a 2a ocorrencia de 'Fonte: E-Agua, 2025.')
foccur = 0
for p in doc.paragraphs:
    if p.text.strip() == "Fonte: E-Agua, 2025.":
        foccur += 1
        if foccur == 2:  # Figura 1.1 (localizacao) -> nosso mapa
            replace_in_paragraph(p, "E-Agua, 2025.", "2S Engenharia e Geotecnologia, 2026.")
print("Fonte fig 1.1 ajustada (ocorrencias E-Agua:", foccur, ")")

# ---------- 11) FIGURAS: trocar blobs image3 (Fig1.1) e image4 (Fig3.1) ----------
def swap_blob(partname_substr, new_path):
    for rel_id, part in doc.part.related_parts.items():
        if partname_substr in str(part.partname):
            with open(new_path, "rb") as f:
                part._blob = f.read()
            return True
    return False
print("swap Fig1.1 (image3 -> Mapa1):", swap_blob("image3.png", MAPA_LOC))
print("swap Fig3.1 (image4 -> Mapa6):", swap_blob("image4.jpeg", MAPA_3D))

# ---------- salvar ----------
doc.save(WORK)
print("SALVO:", WORK)
