# -*- coding: utf-8 -*-
"""
Gera o MODELO de MEMORIAL DESCRITIVO da 2S Engenharia e Geotecnologia.
Projeto Executivo de Engenharia - Rede Coletora de Esgoto (padrao SANEPAR/Acciona).
Documento .docx profissional, com capa, contracapa, sumario (TOC), timbrado 2S e mapas.

Construido com python-docx. Conteudo tecnico generico com {{PLACEHOLDERS}}.
Autor: A2Z / 2S Engenharia - geracao automatica.
"""

import os
import sys
import json
import shutil
import copy
import argparse
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Caminhos / Configuracao
# ----------------------------------------------------------------------------
# A geracao e dirigida por um JSON de configuracao (--config). Quando rodado
# sem config (modo legado / desenvolvimento), cai nos defaults de Amapora.
#
# CFG e o dicionario de configuracao carregado de --config; build_subst_map()
# e build() leem dele. Os caminhos abaixo sao RESOLVIDOS em apply_config().

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(SCRIPT_DIR, "assets")  # template, logos e detalhes embutidos
# Diretorio com os dados/maps do projeto Amapora (modo legado/sample).
BASE = os.environ.get("MEMORIAL_BASE", r"C:\Users\lcabd\jarvis\memorial")
MAPAS = os.path.join(BASE, "mapas")

# Template .docx da 2S (mesmo usado pelo Relatorio Topografico do Nexus):
# ja traz o TIMBRADO (cabecalho/rodape) embutido nos headers/footers, em
# posicao/tamanho corretos (imagens flutuantes ancoradas a pagina, full-bleed),
# alem da configuracao de pagina (pgSz/pgMar) e dos estilos. Partimos DELE.
# Preferimos o template embutido no bundle; caimos no da pasta topografia.
def _first_existing(paths, fallback=None):
    for p in paths:
        if p and os.path.exists(p):
            return p
    return fallback or (paths[0] if paths else None)

TEMPLATE = _first_existing([
    os.path.join(ASSETS, "template_2s.docx"),
    r"D:\PROGRAMAÇÃO\NEXUS\src\modulos\topografia\assets\template_2s.docx",
])
MAPA1 = os.path.join(MAPAS, "Mapa1_Localizacao.png")
MAPA2 = os.path.join(MAPAS, "Mapa2_SubBacia_SB02_Rede.png")
MAPA3 = os.path.join(MAPAS, "Mapa3_Soleiras_SB02.png")
MAPA4 = os.path.join(MAPAS, "Mapa4_Calor_Topografia.png")
MAPA5 = os.path.join(MAPAS, "Mapa5_Interferencias.png")
MAPA6 = os.path.join(MAPAS, "Mapa6_3D_Topografia.png")
MAPA7 = os.path.join(MAPAS, "Mapa7_Sondagem.png")
FLUXOGRAMA = os.path.join(BASE, "Fluxograma_Sistema.png")
FOTOS = os.path.join(BASE, "fotos_campo")
DADOS_JSON = os.path.join(BASE, "dados_amapora.json")
SAIDA = os.path.join(BASE, "Memorial_Descritivo_2S_AMAPORA.docx")

# detalhes-tipo construtivos (CAD) — embutidos no bundle (assets/imageNN.png),
# caindo para a pasta fotos_campo do projeto legado se nao existir no bundle.
def _det(nome_bundle, nome_legado):
    return _first_existing([
        os.path.join(ASSETS, nome_bundle),
        os.path.join(FOTOS, nome_legado),
    ])

DET_TL    = _det("image14.png", "image14.png")  # Terminal de Limpeza
DET_PV_A  = _det("image15.png", "image15.png")  # Poco de Visita Tipo A
DET_PV1000= _det("image16.png", "image16.png")  # PV 1000
DET_PVTR  = _det("image17.png", "image17.png")  # PVTR
DET_TQ    = _det("image18.png", "image18.png")  # Tubo de Queda
DET_LIG1  = _det("image19.png", "image19.png")  # Ligacao Domiciliar tipo 01
DET_LIG2  = _det("image20.png", "image20.png")  # Ligacao Domiciliar tipo 02
DET_LIG3  = _det("image21.png", "image21.png")  # Ligacao Domiciliar tipo 03

# Configuracao do projeto carregada de --config (vazio = modo legado Amapora).
CFG = {}
# Secao 9.5 (soleiras negativas / interferencias) so e emitida quando o config
# aponta esses brutos. Default True = modo legado Amapora (mantem a secao).
TEM_SOLEIRAS = True
TEM_INTERF = True
# Sintese topografica (numero de pontos / cotas) so sai com dados reais de TXT.
# Definido em build_subst_map() pela presenca do bloco 'topografia' nos dados.
TEM_TOPO = True
# GeoJSON da rede projetada (footprint da bacia). Resolvido em apply_config().
REDE_GEOJSON = os.path.join(BASE, "geo_amapora", "rede_trechos.geojson")
# Diretorio de trabalho (escrita do _template_work.docx temporario).
WORKDIR = tempfile.gettempdir()

# ----------------------------------------------------------------------------
# Identidade visual 2S
# ----------------------------------------------------------------------------
COR_GRAFITE = RGBColor(0x30, 0x30, 0x30)   # #303030
COR_VERMELHO = RGBColor(0xA1, 0x13, 0x12)  # #A11312
COR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
COR_CINZA_CLARO = "EDEDED"
COR_CINZA_ZEBRA = "F5F5F5"
FONTE = "Calibri"

COR_FALTA = RGBColor(0xC0, 0x00, 0x00)  # vermelho [A PREENCHER] / [INFORMAR]

# contadores globais de figuras e tabelas (para legendas e listas)
_fig_n = [0]
_tab_n = [0]
FIGURAS = []   # (numero, legenda)
TABELAS = []   # (numero, legenda)


# ----------------------------------------------------------------------------
# DADOS DE AMAPORA + mapa de substituicao de placeholders
# ----------------------------------------------------------------------------
# Placeholders cujo valor ainda NAO existe (nao consta da OSE): vazao/populacao,
# coeficientes, horizonte, indice de atendimento populacional, e os campos do
# fluxograma. Sao substituidos por marcadores VERMELHOS, NUNCA inventados.
FALTA_PREENCHER = "[A PREENCHER]"
FALTA_INFORMAR = "[INFORMAR]"

# tokens marcados em vermelho -> texto do marcador
RED_TOKENS = {}


DADOS_EXTRA_JSON = os.path.join(BASE, "dados_amapora_extra.json")


def _load_dados():
    with open(DADOS_JSON, "r", encoding="utf-8") as f:
        return json.load(f)


def _bacia_area_km2():
    """Area da Bacia 02 = envoltoria convexa (convex hull) dos trechos da rede
    projetada (rede_trechos.geojson, EPSG:31982 em metros), convertida em km2.
    Representa o footprint da rede coletora da sub-bacia."""
    gj_path = REDE_GEOJSON if REDE_GEOJSON else os.path.join(BASE, "geo_amapora", "rede_trechos.geojson")
    try:
        with open(gj_path, "r", encoding="utf-8") as f:
            gj = json.load(f)
    except Exception:
        return 0.0
    pts = []

    def collect(coords):
        if not coords:
            return
        if isinstance(coords[0], (int, float)):
            pts.append((coords[0], coords[1]))
            return
        for c in coords:
            collect(c)

    for feat in gj.get("features", []):
        g = feat.get("geometry") or {}
        collect(g.get("coordinates"))
    if len(pts) < 3:
        return 0.0

    def cross(o, a, b):
        return (a[0] - o[0]) * (b[1] - o[1]) - (a[1] - o[1]) * (b[0] - o[0])

    P = sorted(set(pts))
    lower = []
    for p in P:
        while len(lower) >= 2 and cross(lower[-2], lower[-1], p) <= 0:
            lower.pop()
        lower.append(p)
    upper = []
    for p in reversed(P):
        while len(upper) >= 2 and cross(upper[-2], upper[-1], p) <= 0:
            upper.pop()
        upper.append(p)
    hull = lower[:-1] + upper[:-1]
    A = 0.0
    n = len(hull)
    for i in range(n):
        x1, y1 = hull[i]
        x2, y2 = hull[(i + 1) % n]
        A += x1 * y2 - x2 * y1
    return abs(A) / 2.0 / 1e6


def _load_dados_extra():
    with open(DADOS_EXTRA_JSON, "r", encoding="utf-8") as f:
        return json.load(f)


def _fmt_m(v):
    """Formata numero em metros no padrao pt-BR (milhar com ponto, decimal virgula)."""
    s = "{:,.2f}".format(float(v))
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_int(v):
    s = "{:,d}".format(int(v))
    return s.replace(",", ".")


# dados derivados disponibilizados para build() (tabelas dinamicas)
DADOS = {}        # JSON principal
DADOS_EXTRA = {}  # JSON extra (TQ/degrau, DN, GNSS)
DN_ROWS = []      # linhas reais da tabela de quantitativo de rede [(dn, mat, ext_str)]


def build_subst_map():
    """Monta o dicionario {{PLACEHOLDER}} -> valor real (Amapora) e registra,
    em RED_TOKENS, os marcadores que devem sair em vermelho (dados ausentes)."""
    global DADOS, DADOS_EXTRA, DN_ROWS
    d = _load_dados()
    DADOS = d
    try:
        DADOS_EXTRA = _load_dados_extra()
    except Exception:
        DADOS_EXTRA = {}
    q = d["quantitativos"]
    # TOLERANCIA: soleiras pode nao existir (config sem 'soleiras'). Nesse caso
    # usamos um bloco neutro (zeros) e os placeholders de imoveis/atendimento
    # saem como [A PREENCHER] (ver SOL_AUSENTE abaixo), sem quebrar o memorial.
    sol = d.get("soleiras") or {}
    SOL_AUSENTE = not sol
    sol = {
        "atendidas_positiva": sol.get("atendidas_positiva", 0),
        "nao_atendidas_negativa": sol.get("nao_atendidas_negativa", 0),
        "total_imoveis": sol.get("total_imoveis", 0),
    }
    topo = d.get("topografia") or {}
    ident = d["identificacao"]
    global TEM_TOPO, TEM_SOLEIRAS, TEM_INTERF
    TEM_TOPO = bool(topo.get("total_pontos"))
    # Secao 9.5 e orientada PELOS DADOS (presenca no JSON), nao pelo config —
    # assim o corte vale tambem ao reaproveitar dados.json ja extraido.
    TEM_SOLEIRAS = bool(d.get("soleiras"))
    TEM_INTERF = bool(d.get("interferencias"))

    ext_total = _fmt_m(q["extensao_total_m"])
    prof = q["profundidade_por_faixa_estruturas"]

    m = {}

    # Configuracao do projeto vinda de --config (com fallback p/ Amapora).
    proj = (CFG.get("projeto") or {}) if isinstance(CFG, dict) else {}
    art  = (CFG.get("art") or {}) if isinstance(CFG, dict) else {}

    # --- Identificacao do empreendimento ---
    m["MUNICIPIO"] = proj.get("municipio") or "Amaporã"
    m["UF"] = proj.get("uf") or "PR"
    m["SUBBACIA"] = proj.get("subbacia") or "Bacia 02 (sub-bacia BACIA 02 B)"
    m["EMPREENDIMENTO"] = proj.get("empreendimento") or (
        "Sistema de Esgotamento Sanitário (SES) – Rede Coletora de Esgoto de %s – %s/%s"
        % (m["SUBBACIA"], m["MUNICIPIO"], m["UF"]))
    m["DISTRITOS"] = proj.get("distritos") or ("da área urbana atendida pela %s" % m["SUBBACIA"])

    # --- Investigacao geotecnica (sondagens) ---
    global SOND
    SOND = None
    try:
        import extrair_sondagem as _es
        _sx = CFG.get("sondagem_xlsx") if isinstance(CFG, dict) else None
        SOND = _es.resumo_sondagem(m["MUNICIPIO"], _sx)
    except Exception as _e:
        sys.stderr.write("[sondagem] falha ao extrair: %s\n" % _e)
    if SOND:
        for _k in ("N_FUROS", "METODO_INVEST", "PROF_MIN", "PROF_MED", "PROF_MAX",
                   "SOLO_PRED", "COR_PRED", "SPT_SIM_NAO", "NSPT_MED", "COMPACIDADE_PRED",
                   "NA_SITUACAO", "N_FUROS_NA", "NA_PROF_MED", "DATUM"):
            m["SOND_" + _k] = str(SOND[_k])
        m["SOND_SOLO_L"] = str(SOND["SOLO_PRED"]).lower()
        m["SOND_COR_L"] = str(SOND["COR_PRED"]).lower()

    # --- ART / responsavel tecnico (pre-setado: Zanatta; editavel via config) ---
    eng = art.get("eng_resp") or "RODRIGO ANTONIO ZANATTA"
    m["ENG_RESP"] = eng
    m["CREA"] = art.get("crea") or "SC-0680240/D"
    m["ART"] = art.get("art") or "1720260898051"
    m["ELABORADO"] = art.get("elaborado") or eng
    m["VERIFICADO"] = art.get("verificado") or eng
    m["APROVADO"] = art.get("aprovado") or eng

    # --- Codigo do documento / revisao / data ---
    m["CODIGO_DOC"] = proj.get("codigo_doc") or "MD-RCE-AMAPORA-BACIA02B-R1"
    m["CODIGO_MC"] = proj.get("codigo_mc") or "MC-RCE-AMAPORA-BACIA02B-R1"
    m["REV"] = proj.get("rev") or "R1"
    m["DATA"] = proj.get("data") or "Junho/2026"
    m["DATA_R0"] = proj.get("data_r0") or "Janeiro/2026"
    m["DESCRICAO_REV"] = proj.get("descricao_rev") or (
        "Revisão de projeto executivo com dados de campo consolidados")

    # --- Dados IBGE do municipio (Amapora-PR por padrao) ---
    # Populacao residente / area territorial; editaveis via config.
    POP_MUNI_2022 = int(proj.get("pop_muni") or 4762)
    AREA_MUNI_KM2 = float(proj.get("area_muni") or 384.7)
    m["POP_MUNI"] = _fmt_int(POP_MUNI_2022)
    m["AREA_MUNI"] = "{:.1f}".format(AREA_MUNI_KM2).replace(".", ",")

    # --- Area da Bacia 02 (envoltoria/footprint da rede projetada) ---
    # Convex hull dos trechos da rede (EPSG:31982, metros) -> km2.
    area_bacia_km2 = _bacia_area_km2()
    pct_bacia = area_bacia_km2 / AREA_MUNI_KM2 * 100.0
    # populacao na area da bacia: estimada pelos imoveis atendidos x tamanho
    # medio do domicilio (Censo 2022). PR ~ 2,8 hab/domicilio.
    HAB_POR_DOMIC = 2.8
    pop_bacia = int(round(sol["atendidas_positiva"] * HAB_POR_DOMIC))
    m["AREA_SEDE"] = "{:.2f}".format(area_bacia_km2).replace(".", ",")
    m["PCT_SEDE"] = "{:.2f}%".format(pct_bacia).replace(".", ",")
    m["POP_BACIA"] = FALTA_PREENCHER if SOL_AUSENTE else _fmt_int(pop_bacia)

    # --- Base topografica (dados reais GNSS) ---
    m["REF_ALTIMETRICO"] = "SIRGAS 2000 / UTM 22S (EPSG:31982)"
    gnss = DADOS_EXTRA.get("3_precisoes_gnss", {}) if DADOS_EXTRA else {}
    if gnss:
        prec_h = gnss.get("precisao_planimetrica_m_HRMS", {}).get("tipico_mediana", 0.008)
        prec_v = gnss.get("precisao_altimetrica_m_VRMS", {}).get("tipico_mediana", 0.014)
        n_pts = gnss.get("n_pontos", 10035)
        pdop_med = gnss.get("PDOP", {}).get("medio", 0.69)
        metodo_raw = gnss.get("metodo_RTK", "PPP-RTK (GEO PPP)")
        # usa apenas a designacao do metodo (descarta detalhes apos " - ")
        metodo = metodo_raw.split(" - ")[0].strip() if metodo_raw else "PPP-RTK (GEO PPP)"
        _ph = "{:.3f}".format(prec_h).replace(".", ",")
        _pv = "{:.3f}".format(prec_v).replace(".", ",")
        m["PREC_PLANI"] = "~%s m (HRMS), obtida pelo método %s" % (_ph, metodo)
        m["PREC_ALTI"] = "~%s m (VRMS)" % _pv
        m["GNSS_N_PONTOS"] = _fmt_int(n_pts)
        m["GNSS_METODO"] = metodo
        m["GNSS_PDOP"] = "{:.2f}".format(pdop_med).replace(".", ",")
        m["GNSS_PREC_PLANI"] = "{:.3f}".format(prec_h).replace(".", ",")
        m["GNSS_PREC_ALTI"] = "{:.3f}".format(prec_v).replace(".", ",")
    else:
        m["PREC_PLANI"] = FALTA_PREENCHER
        m["PREC_ALTI"] = FALTA_PREENCHER
        m["GNSS_N_PONTOS"] = "10.035"
        m["GNSS_METODO"] = "PPP-RTK (GEO PPP)"
        m["GNSS_PDOP"] = "0,69"
        m["GNSS_PREC_PLANI"] = "0,008"
        m["GNSS_PREC_ALTI"] = "0,014"
    # pontos de apoio / RRNN: nao temos a quantidade -> vermelho
    m["QTD_APOIO"] = FALTA_PREENCHER

    # --- Sintese do levantamento topografico (dados REAIS do TXT) ---
    if TEM_TOPO:
        _zmin = topo.get("z_min"); _zmax = topo.get("z_max"); _zmed = topo.get("z_media")
        m["TOPO_N_PONTOS"] = _fmt_int(topo.get("total_pontos", 0))
        m["TOPO_Z_MIN"] = (_fmt_m(_zmin) if _zmin is not None else FALTA_PREENCHER)
        m["TOPO_Z_MAX"] = (_fmt_m(_zmax) if _zmax is not None else FALTA_PREENCHER)
        m["TOPO_Z_MED"] = (_fmt_m(_zmed) if _zmed is not None else FALTA_PREENCHER)
        if _zmin is not None and _zmax is not None:
            m["TOPO_DESNIVEL"] = _fmt_m(_zmax - _zmin)
        else:
            m["TOPO_DESNIVEL"] = FALTA_PREENCHER
    else:
        for _k in ("TOPO_N_PONTOS", "TOPO_Z_MIN", "TOPO_Z_MAX", "TOPO_Z_MED", "TOPO_DESNIVEL"):
            m[_k] = FALTA_PREENCHER

    # --- Criterios / parametros (vazoes e coeficientes: NAO existem) -> vermelho ---
    m["QPERC"] = FALTA_PREENCHER
    m["K1"] = FALTA_PREENCHER
    m["K2"] = FALTA_PREENCHER
    m["K3"] = FALTA_PREENCHER
    m["C_RETORNO"] = FALTA_PREENCHER
    m["TX_INFILTRACAO"] = FALTA_PREENCHER
    m["DECL_MIN"] = FALTA_PREENCHER
    m["DECL_MIN_PARAM"] = FALTA_PREENCHER
    m["RECOB_MIN"] = "0,95"
    m["TAXA_LINEAR"] = FALTA_PREENCHER
    # defaults dos parametros de rede (sobrescritos pelo modelo, se apontado)
    m["MANNING"] = "0,010"
    m["MATERIAL_REDE"] = "PVC"
    m["DN_MIN"] = "150"

    # --- Modelo hidraulico SewerGEMS (.sqlite): preenche os parametros reais ---
    # da secao 9.6 quando o usuario aponta o modelo. Sem modelo, mantem os
    # defaults/[A PREENCHER] acima.
    mod = DADOS_EXTRA.get("4_modelo_hidraulico") if DADOS_EXTRA else None
    if mod:
        _decl = mod.get("declividade_min_m_m")
        if _decl is not None:
            m["DECL_MIN_PARAM"] = ("{:.4f}".format(_decl).replace(".", ","))
            m["DECL_MIN"] = m["DECL_MIN_PARAM"]
        if mod.get("recobrimento_min_m") is not None:
            m["RECOB_MIN"] = "{:.2f}".format(mod["recobrimento_min_m"]).replace(".", ",")
        if mod.get("manning"):
            m["MANNING"] = "{:.3f}".format(mod["manning"]).replace(".", ",")
        if mod.get("material"):
            m["MATERIAL_REDE"] = str(mod["material"])
        if mod.get("dn_min_mm"):
            m["DN_MIN"] = str(mod["dn_min_mm"])

    # --- Horizonte de projeto / populacao / vazoes (NAO existem) -> vermelho ---
    m["HORIZONTE_ANOS"] = FALTA_PREENCHER
    m["ANO_INI"] = FALTA_PREENCHER
    m["ANO_FIM"] = FALTA_PREENCHER
    m["POP_INI"] = FALTA_PREENCHER
    m["EXT_INI"] = FALTA_PREENCHER
    m["QINF_INI"] = FALTA_PREENCHER
    m["QMED_INI"] = FALTA_PREENCHER
    m["QMAX_INI"] = FALTA_PREENCHER
    m["POP_FIM"] = FALTA_PREENCHER
    m["EXT_FIM"] = FALTA_PREENCHER
    m["QINF_FIM"] = FALTA_PREENCHER
    m["QMED_FIM"] = FALTA_PREENCHER
    m["QMAX_FIM"] = FALTA_PREENCHER
    m["OBS_VAZOES_CONCENTRADAS"] = FALTA_PREENCHER

    # --- Fluxograma (EEE / ETE / corpo receptor: nao informados) -> vermelho ---
    m["EEE"] = FALTA_INFORMAR
    m["ETE"] = FALTA_INFORMAR
    m["CORPO_RECEPTOR"] = FALTA_INFORMAR

    # --- Soleiras negativas (dados reais - Bacia 02) ---
    if SOL_AUSENTE:
        # sem cadastro de soleiras: marca imoveis/atendimento como a preencher
        m["IMOVEIS_ATEND"] = FALTA_PREENCHER
        m["IMOVEIS_NAO_ATEND"] = FALTA_PREENCHER
        m["IMOVEIS_TOTAL"] = FALTA_PREENCHER
        m["PCT_ATEND"] = FALTA_PREENCHER
        m["PCT_NAO_ATEND"] = FALTA_PREENCHER
    else:
        m["IMOVEIS_ATEND"] = _fmt_int(sol["atendidas_positiva"])
        m["IMOVEIS_NAO_ATEND"] = _fmt_int(sol["nao_atendidas_negativa"])
        m["IMOVEIS_TOTAL"] = _fmt_int(sol["total_imoveis"])
        # indice de atendimento das soleiras: 1 casa decimal (ex.: 98,9%)
        _tot = sol["total_imoveis"] or 0
        pct_at = (sol["atendidas_positiva"] / _tot * 100.0) if _tot else 0.0
        m["PCT_ATEND"] = "{:.1f}%".format(pct_at).replace(".", ",")
        m["PCT_NAO_ATEND"] = "{:.1f}%".format(100 - pct_at).replace(".", ",")
    # Total de soleiras (campo unico p/ a tabela de dispositivos): real quando ha
    # cadastro, senao [A PREENCHER] (preenchimento manual pelas pranchas).
    m["QTD_SOLEIRAS"] = (FALTA_PREENCHER if SOL_AUSENTE
                         else _fmt_int(sol.get("total_imoveis") or 0))

    # --- Interferencias com redes existentes (dados reais) ---
    interf = d.get("interferencias", {})
    m["INTERF_AGUA"] = _fmt_int(interf.get("agua_trechos", 0)) if interf else FALTA_PREENCHER
    m["INTERF_DREN"] = _fmt_int(interf.get("drenagem_trechos", 0)) if interf else FALTA_PREENCHER

    # --- Quantitativos de rede (dados reais) ---
    # Tabela de rede montada SO com os DN realmente presentes (dados extra),
    # caindo para o DN unico do JSON principal se o extra nao existir.
    m["EXT_TOTAL"] = ext_total
    DN_ROWS.clear()
    dn_src = (DADOS_EXTRA.get("2_dn_rede", {}).get("dn_extensao")
              if DADOS_EXTRA else None)
    soma_dn = 0.0
    if dn_src:
        for r in dn_src:
            dn = str(r.get("dn", "")).replace("DN", "").strip()
            mat = r.get("material", "PVC")
            ext = float(r.get("extensao_m", 0))
            soma_dn += ext
            DN_ROWS.append(("DN %s mm" % dn, mat, _fmt_m(ext)))
    else:
        # fallback: DN unico do JSON principal
        for dn_label, ext in q.get("extensao_por_DN_m", {}).items():
            dn = dn_label.replace("DN", "").replace("mm - PVC", "").replace("mm", "").strip()
            mat = "PVC"
            soma_dn += float(ext)
            DN_ROWS.append(("DN %s mm" % dn, mat, _fmt_m(ext)))
    # total da tabela de rede = soma dos DN presentes
    m["EXT_TOTAL_REDE"] = _fmt_m(soma_dn if soma_dn else q["extensao_total_m"])

    m["QTD_PV"] = _fmt_int(q["n_PV"])
    m["QTD_TL"] = _fmt_int(q["n_TL"])
    # Tubos de queda e degraus: dados reais do JSON extra
    tqd = DADOS_EXTRA.get("1_tubos_de_queda_e_degrau", {}) if DADOS_EXTRA else {}
    n_tq = tqd.get("tubos_de_queda_TQ", {}).get("quantidade")
    soma_tq = tqd.get("tubos_de_queda_TQ", {}).get("soma_alturas_m")
    n_deg = tqd.get("degraus", {}).get("quantidade")
    soma_deg = tqd.get("degraus", {}).get("soma_alturas_m")
    m["QTD_TQ"] = _fmt_int(n_tq) if n_tq is not None else FALTA_PREENCHER
    m["QTD_DEGRAU"] = _fmt_int(n_deg) if n_deg is not None else FALTA_PREENCHER
    m["TQ_SOMA_ALT"] = _fmt_m(soma_tq) if soma_tq is not None else FALTA_PREENCHER
    m["DEGRAU_SOMA_ALT"] = _fmt_m(soma_deg) if soma_deg is not None else FALTA_PREENCHER
    # TIL: vinculado ao numero de ligacoes prediais. Provisoriamente igual ao
    # numero de imoveis atendidos da Bacia 02 (457) - a confirmar.
    n_til = sol["atendidas_positiva"]
    m["QTD_TIL"] = FALTA_PREENCHER if SOL_AUSENTE else _fmt_int(n_til)
    m["QTD_TIL_NOTA"] = "(= nº de ligações prediais — confirmar)"
    # Total de dispositivos contados (PV + TL + TQ + degraus + TIL)
    _disp = q["n_PV"] + q["n_TL"] + n_til
    if n_tq is not None:
        _disp += n_tq
    if n_deg is not None:
        _disp += n_deg
    m["QTD_DISP_TOTAL"] = _fmt_int(_disp)

    # --- Faixas de profundidade: o JSON traz CONTAGEM DE ESTRUTURAS, nao extensao.
    # A extensao por faixa virá da planilha de calculo; aqui informamos o nº de
    # estruturas (o rotulo da tabela e adaptado para "Nº de estruturas").
    m["EXT_PROF_125"] = _fmt_int(prof["ate_1_25"])
    m["EXT_PROF_200"] = _fmt_int(prof["1_25_a_2_00"])
    m["EXT_PROF_300"] = _fmt_int(prof["2_00_a_3_00"])
    m["EXT_PROF_400"] = _fmt_int(prof["3_00_a_4_00"])
    m["EXT_PROF_500"] = _fmt_int(prof["acima_4_00"])

    # --- Metodo executivo (dados reais) ---
    me = q["metodo_extensao"]
    mnd = me["MND_m"]
    vca = me["VCA_m"]
    tot = mnd + vca
    m["EXT_MND"] = _fmt_m(mnd)
    m["EXT_VCA"] = _fmt_m(vca)
    m["PCT_MND"] = "{:.2f}%".format(mnd / tot * 100).replace(".", ",")
    m["PCT_VCA"] = "{:.2f}%".format(vca / tot * 100).replace(".", ",")

    # --- OVERRIDE de quantitativo de rede (config: quantitativos_rede) ---------
    # Fonte autoritativa da extensao/DN (ex.: tabela consolidada do projeto),
    # usada quando o somatorio das OSE nao bate (multi-bacia sem RESUMO unico).
    # Estrutura: {"extensao_total_m": 35215.99,
    #             "dn": [{"dn":"150","material":"PVC","ext":34916.39}, ...]}
    qr = (CFG.get("quantitativos_rede") if isinstance(CFG, dict) else None)
    if qr:
        tot_ov = qr.get("extensao_total_m")
        dn_ov = qr.get("dn") or []
        if dn_ov:
            DN_ROWS.clear()
            soma_ov = 0.0
            for r in dn_ov:
                dn = str(r.get("dn", "")).replace("DN", "").replace("mm", "").strip()
                mat = r.get("material", "PVC")
                ext = float(r.get("ext", r.get("extensao_m", 0)) or 0)
                soma_ov += ext
                DN_ROWS.append(("DN %s mm" % dn, mat, _fmt_m(ext)))
            if tot_ov is None:
                tot_ov = soma_ov
            m["EXT_TOTAL_REDE"] = _fmt_m(soma_ov)
        if tot_ov is not None:
            m["EXT_TOTAL"] = _fmt_m(tot_ov)
        # Metodo executivo: VCA = extensao das planilhas onde prof > 3,00 m
        # (somada de TODAS as abas OSE); MND = TOTAL(imagem) - VCA. Assim o
        # metodo fecha exatamente no quantitativo informado.
        vca_pl = q.get("metodo_extensao", {}).get("VCA_planilhas_m")
        if tot_ov is not None and vca_pl is not None:
            vca_pl = float(vca_pl)
            mnd_pl = max(tot_ov - vca_pl, 0.0)
            m["EXT_VCA"] = _fmt_m(vca_pl)
            m["EXT_MND"] = _fmt_m(mnd_pl)
            m["PCT_VCA"] = "{:.2f}%".format(vca_pl / tot_ov * 100).replace(".", ",")
            m["PCT_MND"] = "{:.2f}%".format(mnd_pl / tot_ov * 100).replace(".", ",")
        # PV/TL e faixas de profundidade: com override multi-bacia o RESUMO nao
        # cobre todas as bacias (metodo antigo) -> marca [A PREENCHER] p/ contagem
        # manual pelas pranchas do CAD. TQ/degrau (lidos de cada aba) seguem reais.
        m["QTD_PV"] = FALTA_PREENCHER
        m["QTD_TL"] = FALTA_PREENCHER
        m["QTD_TIL"] = FALTA_PREENCHER
        m["QTD_DISP_TOTAL"] = FALTA_PREENCHER
        m["EXT_PROF_125"] = FALTA_PREENCHER
        m["EXT_PROF_200"] = FALTA_PREENCHER
        m["EXT_PROF_300"] = FALTA_PREENCHER
        m["EXT_PROF_400"] = FALTA_PREENCHER
        m["EXT_PROF_500"] = FALTA_PREENCHER

    # registra marcadores vermelhos
    for k, v in m.items():
        if v in (FALTA_PREENCHER, FALTA_INFORMAR):
            RED_TOKENS[v] = True

    return m


SUBST = {}  # preenchido em build()


def _substitute_runs_in_paragraph(p):
    """Substitui {{TOKENS}} nos runs de um paragrafo. Mantem a formatacao do
    run; tokens ausentes (FALTA_*) recebem fonte vermelha.

    Estrategia: concatena o texto do paragrafo, faz a substituicao gerando
    segmentos (texto, vermelho?), e reescreve preservando as propriedades do
    PRIMEIRO run como base de formatacao."""
    import re
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return
    if not p.runs:
        return

    # Se o paragrafo contiver CAMPOS do Word (ex.: SEQ nas legendas Caption),
    # NAO podemos reconstruir os runs (isso apagaria o campo). Nesse caso fazemos
    # substituicao token-a-token DENTRO de cada run de texto, preservando os
    # runs de campo. Cada token (ex.: {{SUBBACIA}}) e adicionado como run unico
    # em add_caption, entao esta abordagem cobre 100% das legendas.
    has_field = p._p.find(qn('w:r') + '/' + qn('w:fldChar')) is not None
    if not has_field:
        for el in p._p.iter(qn('w:fldChar')):
            has_field = True
            break
    if has_field:
        for r in list(p.runs):
            if "{{" not in r.text:
                continue
            def _rep(mt):
                return SUBST.get(mt.group(1), mt.group(0))
            new_text = re.sub(r"\{\{([A-Z0-9_]+)\}\}", _rep, r.text)
            # se virou um marcador vermelho, pinta o run
            if new_text in (FALTA_PREENCHER, FALTA_INFORMAR):
                r.font.color.rgb = COR_FALTA
                r.bold = True
            r.text = new_text
        return

    base_run = p.runs[0]
    base_rpr = base_run._r.find(qn('w:rPr'))

    # quebra em segmentos preservando deteccao de vermelho
    segments = []  # (texto, is_red)
    pos = 0
    for mt in re.finditer(r"\{\{([A-Z0-9_]+)\}\}", full):
        if mt.start() > pos:
            segments.append((full[pos:mt.start()], False))
        key = mt.group(1)
        val = SUBST.get(key, mt.group(0))
        is_red = val in (FALTA_PREENCHER, FALTA_INFORMAR)
        segments.append((val, is_red))
        pos = mt.end()
    if pos < len(full):
        segments.append((full[pos:], False))

    # remove todos os runs existentes
    for r in list(p.runs):
        r._r.getparent().remove(r._r)

    # recria runs com a formatacao base
    for text, is_red in segments:
        if text == "":
            continue
        nr = p.add_run(text)
        if base_rpr is not None:
            new_rpr = copy.deepcopy(base_rpr)
            # remove cor existente para nao conflitar
            for c in new_rpr.findall(qn('w:color')):
                new_rpr.remove(c)
            old = nr._r.find(qn('w:rPr'))
            if old is not None:
                nr._r.remove(old)
            nr._r.insert(0, new_rpr)
        if is_red:
            nr.font.color.rgb = COR_FALTA
            nr.bold = True


def apply_substitutions(doc):
    """Percorre todos os paragrafos do corpo e das tabelas substituindo tokens."""
    for p in doc.paragraphs:
        _substitute_runs_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _substitute_runs_in_paragraph(p)


# ----------------------------------------------------------------------------
# Helpers de baixo nivel (XML)
# ----------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        node = OxmlElement('w:' + tag)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="BFBFBF", size=6):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(size))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)
        borders.append(e)
    tblPr.append(borders)


def repeat_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    th.set(qn('w:val'), 'true')
    trPr.append(th)


def add_field(paragraph, instr_text):
    """Insere um campo de Word (ex.: PAGE, TOC) num paragrafo."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    run._r.append(instr)

    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldSep)

    # texto provisorio (atualizado ao abrir/F9)
    t = OxmlElement('w:t')
    t.text = ""
    run._r.append(t)

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldEnd)
    return run


def add_toc(paragraph, instr):
    """Insere um campo composto (TOC / lista de figuras/tabelas) com placeholder."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)
    instr_el = OxmlElement('w:instrText')
    instr_el.set(qn('xml:space'), 'preserve')
    instr_el.text = instr
    run._r.append(instr_el)
    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldSep)

    run2 = paragraph.add_run("  [Atualize este campo no Word: clique e pressione F9]")
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    run3 = paragraph.add_run()
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldEnd)


def _add_seq_field(paragraph, label):
    """Insere um campo SEQ ( SEQ <label> \\* ARABIC ) num paragrafo, devolvendo
    o run que o contem. Esse campo e o que a Lista de Figuras/Tabelas (TOC \\c
    "<label>") usa para identificar e numerar as legendas no F9."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' SEQ %s \\* ARABIC ' % label
    run._r.append(instr)
    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldSep)
    t = OxmlElement('w:t')
    t.text = "0"  # texto provisorio; o F9 substitui pelo numero do SEQ
    run._r.append(t)
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldEnd)
    return run


def ensure_caption_style(doc):
    """Garante o estilo 'Caption' (Legenda) no documento. E ESSE estilo que o
    campo TOC \\c usa para coletar as legendas. python-docx mapeia 'Caption'."""
    if getattr(doc, "_caption_style", None) is not None:
        return doc._caption_style
    styles = doc.styles
    st = None
    for s in styles:
        if s.name in ("Caption", "Legenda"):
            st = s
            break
    if st is None:
        st = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles['Normal']
    st.font.name = FONTE
    st.font.size = Pt(9.5)
    st.font.bold = True
    st.font.italic = False
    st.font.color.rgb = COR_GRAFITE
    pf = st.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(4)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.0
    doc._caption_style = st
    return st


def add_caption(doc, label, numero, texto, align=None):
    """Cria um paragrafo de legenda no estilo 'Caption' com o padrao
    "<label> <N> – <texto>", usando NUMERO LITERAL (sem campo SEQ). Assim a
    numeracao ja sai correta ao abrir o documento, sem depender de F9, e as
    Listas de Figuras/Tabelas sao montadas literalmente pelo gerador."""
    ensure_caption_style(doc)
    p = doc.add_paragraph(style="Caption")
    if align is not None:
        p.alignment = align
    r0 = p.add_run("%s %d – %s" % (label, numero, texto))
    r0.bold = True
    r0.font.size = Pt(9.5)
    r0.font.name = FONTE
    r0.font.color.rgb = COR_GRAFITE
    return p


def _insert_list_paragraph_after(doc, anchor_p, label, numero, texto):
    """Cria um paragrafo de item de lista ("<label> N – titulo  ....") e o
    insere LOGO APOS o paragrafo `anchor_p`. Devolve o novo paragrafo (para
    encadear). Usa tab com preenchimento por pontos ate a margem direita."""
    new_p = OxmlElement('w:p')
    anchor_p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, anchor_p._parent)
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    # tab de pontos perto da margem direita (~16 cm)
    pPr = new_p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(int(Cm(16).emu / 635)))  # twips
    tabs.append(tab)
    pPr.append(tabs)
    r = p.add_run("%s %d – %s" % (label, numero, texto))
    r.font.size = Pt(10.5)
    r.font.name = FONTE
    r.font.color.rgb = COR_GRAFITE
    return p


def page_width_emu(section):
    """Largura util da pagina (descontadas as margens), em EMU."""
    return section.page_width - section.left_margin - section.right_margin


# ----------------------------------------------------------------------------
# Helpers de conteudo
# ----------------------------------------------------------------------------
def _ensure_bullet_numbering(doc):
    """Garante uma definicao de numeracao com marcador (bullet) no
    numbering.xml e devolve o numId correspondente. Cria abstractNum + num
    se ainda nao existirem (o template nao traz lista com marcador)."""
    if getattr(doc, "_bullet_numId", None) is not None:
        return doc._bullet_numId
    numbering = doc.part.numbering_part.element
    # ids livres
    abs_ids = [int(a.get(qn('w:abstractNumId')))
               for a in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(n.get(qn('w:numId')))
               for n in numbering.findall(qn('w:num'))]
    abs_id = (max(abs_ids) + 1) if abs_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    # abstractNum com 2 niveis de bullet
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abs_id))
    for lvl, (char, font, left, hang) in enumerate([
            ("", "Symbol", 360, 360),     # nivel 0: bolinha cheia
            ("–", FONTE, 720, 360)]):       # nivel 1: travessao
        l = OxmlElement('w:lvl'); l.set(qn('w:ilvl'), str(lvl))
        s = OxmlElement('w:start'); s.set(qn('w:val'), '1'); l.append(s)
        nf = OxmlElement('w:numFmt'); nf.set(qn('w:val'), 'bullet'); l.append(nf)
        lt = OxmlElement('w:lvlText'); lt.set(qn('w:val'), char); l.append(lt)
        lj = OxmlElement('w:lvlJc'); lj.set(qn('w:val'), 'left'); l.append(lj)
        ppr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left)); ind.set(qn('w:hanging'), str(hang))
        ppr.append(ind); l.append(ppr)
        rpr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font)
        rpr.append(rf); l.append(rpr)
        abstract.append(l)
    # abstractNum deve vir antes dos <w:num>
    first_num = numbering.find(qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering.append(abstract)

    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(num_id))
    a = OxmlElement('w:abstractNumId'); a.set(qn('w:val'), str(abs_id))
    num.append(a)
    numbering.append(num)

    doc._bullet_numId = num_id
    return num_id


def ensure_list_styles(doc):
    """Cria os estilos 'List Bullet' e 'List Bullet 2' se o template nao os
    tiver (ele so traz 'List Paragraph'). Ligados a numeracao de marcador."""
    num_id = _ensure_bullet_numbering(doc)
    styles = doc.styles
    existing = {s.name for s in styles}
    for name, ilvl in (("List Bullet", 0), ("List Bullet 2", 1)):
        if name in existing:
            continue
        st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = styles['Normal']
        st.font.name = FONTE
        st.font.size = Pt(11)
        ppr = st.element.get_or_add_pPr()
        numpr = OxmlElement('w:numPr')
        ilvl_el = OxmlElement('w:ilvl'); ilvl_el.set(qn('w:val'), str(ilvl))
        numid_el = OxmlElement('w:numId'); numid_el.set(qn('w:val'), str(num_id))
        numpr.append(ilvl_el); numpr.append(numid_el)
        ppr.append(numpr)


def style_base(doc):
    """Configura o estilo Normal e os Headings."""
    ensure_list_styles(doc)
    normal = doc.styles['Normal']
    normal.font.name = FONTE
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles['Heading 1']
    h1.font.name = FONTE
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = COR_VERMELHO
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles['Heading 2']
    h2.font.name = FONTE
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COR_GRAFITE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles['Heading 3']
    h3.font.name = FONTE
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = COR_GRAFITE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def add_para(doc, text, size=11, bold=False, italic=False, align=None,
             color=None, space_after=None, space_before=None, line=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = FONTE
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if line is not None:
        p.paragraph_format.line_spacing = line
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = FONTE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    return p


def add_figure(doc, img_path, legenda, fonte="Fonte: 2S Engenharia, 2026."):
    # TOLERANCIA A MAPA FALTANTE: se o PNG da figura nao existir (ex.: o mapa
    # nao foi gerado por falta de dados/interferencias), PULAMOS a figura
    # graciosamente — o memorial sai mesmo assim. Como o contador (_fig_n) so
    # avanca quando a figura entra de fato, a numeracao das figuras seguintes e
    # a Lista de Figuras (montada de FIGURAS) permanecem corretas, sem buracos.
    if not img_path or not os.path.exists(img_path):
        sys.stderr.write(
            "[memorial] AVISO: figura ausente, pulada -> %s (legenda: %s)\n"
            % (img_path, legenda))
        return None

    _fig_n[0] += 1
    n = _fig_n[0]
    FIGURAS.append((n, legenda))

    # imagem centralizada, ajustada a largura util
    section = doc.sections[-1]
    avail = page_width_emu(section)
    # nao estourar; mapas tem proporcao ~1.414 (paisagem) -> usa 95% da largura
    width = int(avail * 0.92)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Emu(width))

    # legenda no estilo Caption com NUMERO LITERAL (popula a Lista de Figuras)
    add_caption(doc, "Figura", n, legenda, align=WD_ALIGN_PARAGRAPH.CENTER)

    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(10)
    rs = src.add_run(fonte)
    rs.italic = True
    rs.font.size = Pt(8.5)
    rs.font.name = FONTE
    rs.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    return n


def add_table_caption(doc, legenda):
    _tab_n[0] += 1
    n = _tab_n[0]
    TABELAS.append((n, legenda))
    # legenda no estilo Caption com NUMERO LITERAL (popula a Lista de Tabelas)
    add_caption(doc, "Tabela", n, legenda, align=WD_ALIGN_PARAGRAPH.LEFT)
    return n


def add_table_source(doc, fonte="Fonte: 2S Engenharia, 2026."):
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.LEFT
    src.paragraph_format.space_after = Pt(10)
    rs = src.add_run(fonte)
    rs.italic = True
    rs.font.size = Pt(8.5)
    rs.font.name = FONTE
    rs.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def make_table(doc, headers, rows, header_color=None, col_widths=None,
               first_col_left=False, zebra=True):
    """Cria tabela com cabecalho colorido (grafite por padrao), zebra e bordas."""
    if header_color is None:
        header_color = "303030"
    ncols = len(headers)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    hdr = table.rows[0]
    repeat_header_row(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = FONTE
        run.font.color.rgb = COR_BRANCO

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra and ridx % 2 == 1:
                set_cell_bg(cell, COR_CINZA_ZEBRA)
            para = cell.paragraphs[0]
            if first_col_left and i == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = FONTE

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ----------------------------------------------------------------------------
# Capa, contracapa, cabecalho/rodape
# ----------------------------------------------------------------------------
def add_header_footer(section, with_page_number=True):
    """NO-OP intencional: o timbrado (cabecalho/rodape + numero de pagina) JA
    vem do TEMPLATE da 2S (template_2s.docx), exatamente como no Relatorio
    Topografico do Nexus.

    No template, o cabecalho e o rodape sao IMAGENS FLUTUANTES ancoradas a
    PAGINA (full-bleed), e o numero de pagina e uma caixa de texto no rodape.

    NAO tocamos no header/footer aqui de proposito:
    - a 1a secao mantem a referencia de header/footer do TEMPLATE;
    - cada nova secao criada por add_section nasce SEM headerReference, o que
      no Word significa "vinculado ao anterior" (= mesmo timbrado).
    Mexer em is_linked_to_previous na 1a secao APAGARIA a referencia do
    template (nao ha secao anterior), fazendo o timbrado sumir. Por isso,
    no-op total. Qualquer insercao de imagem inline aqui era justamente o que
    deslocava/aumentava o timbrado.
    """
    return


def build_cover(doc):
    """Capa com timbrado + numeracao (1a secao)."""
    section = doc.sections[0]
    add_header_footer(section, with_page_number=True)
    avail = page_width_emu(section)

    # faixa de respiro abaixo do cabecalho
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(36)

    # linha vermelha decorativa
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = line.add_run("—" * 18)
    rl.font.color.rgb = COR_VERMELHO
    rl.font.size = Pt(10)
    line.paragraph_format.space_after = Pt(28)

    # tipo de documento
    add_para(doc, "PROJETO EXECUTIVO DE ENGENHARIA", size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COR_GRAFITE, space_after=6)

    # titulo grande
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = t.add_run("MEMORIAL DESCRITIVO")
    rt.bold = True
    rt.font.size = Pt(40)
    rt.font.name = FONTE
    rt.font.color.rgb = COR_VERMELHO
    t.paragraph_format.space_after = Pt(8)

    add_para(doc, "Rede Coletora de Esgoto Sanitário", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COR_GRAFITE, space_after=30)

    # bloco de identificacao do empreendimento
    add_para(doc, "{{MUNICIPIO}} – {{UF}}", size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=COR_GRAFITE, space_after=2)
    add_para(doc, "Sub-Bacia {{SUBBACIA}}  –  Sistema de Esgotamento Sanitário (SES)",
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=COR_GRAFITE, space_after=36)

    # linha vermelha decorativa
    line2 = doc.add_paragraph()
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl2 = line2.add_run("—" * 18)
    rl2.font.color.rgb = COR_VERMELHO
    rl2.font.size = Pt(10)
    line2.paragraph_format.space_after = Pt(30)

    # contratante / consorcio
    add_para(doc, "CONTRATANTE", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COR_VERMELHO, space_after=0)
    add_para(doc, "ACCIONA", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COR_GRAFITE, space_after=12)

    add_para(doc, "PROJETISTA", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=COR_VERMELHO, space_after=0)
    add_para(doc, "Consórcio E-Água e 2S Engenharia e Geotecnologia",
             size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=COR_GRAFITE,
             space_after=36)

    # codigo / revisao / data
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(10)
    ri = info.add_run("Documento nº {{CODIGO_DOC}}     |     Revisão {{REV}}     |     {{DATA}}")
    ri.font.size = Pt(11)
    ri.bold = True
    ri.font.name = FONTE
    ri.font.color.rgb = COR_GRAFITE


def new_section(doc, with_page_number=True):
    """Inicia nova secao em nova pagina.

    NAO redefine a geometria de pagina: a nova secao HERDA do template
    (pgSz A4, pgMar top=4,0 cm / left=right=3,0 cm / bottom=2,5 cm, header=
    footer=1,25 cm) e tambem o timbrado (link to previous). Forcar margens
    proprias aqui era o que brigava com o timbrado flutuante do template.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(section, with_page_number=with_page_number)
    return section


def new_section_landscape(doc):
    """Inicia nova secao em PAISAGEM (A4 deitado), herdando o timbrado do
    template (link to previous). Troca largura<->altura e ajusta a orientacao
    apenas nesta secao."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(section)
    section.orientation = WD_ORIENT.LANDSCAPE
    w, h = section.page_width, section.page_height
    section.page_width = h
    section.page_height = w
    return section


def new_section_portrait(doc):
    """Volta a RETRATO (A4 em pe) numa nova secao, herdando o timbrado."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_header_footer(section)
    section.orientation = WD_ORIENT.PORTRAIT
    w, h = section.page_width, section.page_height
    if w > h:  # garante retrato
        section.page_width = h
        section.page_height = w
    return section


def add_detalhe(doc, img_path, legenda, fonte="Fonte: 2S Engenharia."):
    """Insere um detalhe-tipo construtivo (desenho CAD) centralizado, com
    legenda de figura. Limita a altura para nao estourar a pagina."""
    _fig_n[0] += 1
    n = _fig_n[0]
    FIGURAS.append((n, legenda))

    section = doc.sections[-1]
    avail = page_width_emu(section)
    width = int(avail * 0.82)
    # altura util da pagina (descontadas margens) para nao estourar
    avail_h = section.page_height - section.top_margin - section.bottom_margin
    max_h = int(avail_h * 0.62)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    pic = run.add_picture(img_path, width=Emu(width))
    if pic.height > max_h:
        ratio = max_h / pic.height
        pic.height = Emu(int(pic.height * ratio))
        pic.width = Emu(int(pic.width * ratio))

    # legenda no estilo Caption com NUMERO LITERAL (popula a Lista de Figuras)
    add_caption(doc, "Figura", n, legenda, align=WD_ALIGN_PARAGRAPH.CENTER)

    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_after = Pt(10)
    rs = src.add_run(fonte)
    rs.italic = True
    rs.font.size = Pt(8.5)
    rs.font.name = FONTE
    rs.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    return n


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


# ============================================================================
# CONSTRUCAO DO DOCUMENTO
# ============================================================================
def _set_update_fields_on_open(doc):
    """Adiciona <w:updateFields w:val="true"/> ao settings.xml para que o Word
    atualize os campos automaticos (SUMARIO, numero de pagina) ao ABRIR o
    documento, dispensando o F9. (As listas de figuras/tabelas ja sao
    literais, mas isso garante o sumario de secoes preenchido.)"""
    try:
        settings = doc.settings.element
    except Exception:
        return
    # remove existente
    for el in settings.findall(qn('w:updateFields')):
        settings.remove(el)
    uf = OxmlElement('w:updateFields')
    uf.set(qn('w:val'), 'true')
    settings.insert(0, uf)


def open_template_clean():
    """Abre uma copia do template_2s.docx e LIMPA o corpo, preservando os
    headers/footers (timbrado), os estilos e a configuracao de pagina
    (sectPr final: pgSz/pgMar + headerReference/footerReference).

    Assim o memorial nasce com o MESMO timbrado do Relatorio Topografico.
    """
    if not os.path.exists(TEMPLATE):
        raise FileNotFoundError(
            "Template da 2S nao encontrado: %s\n"
            "E o mesmo usado pelo Relatorio Topografico do Nexus." % TEMPLATE)
    tmp = os.path.join(WORKDIR, "_template_work.docx")
    shutil.copyfile(TEMPLATE, tmp)
    doc = Document(tmp)

    # Remove todo o conteudo do corpo (paragrafos e tabelas do relatorio
    # topografico), MAS preserva o sectPr final do body -- que carrega as
    # referencias de header/footer (timbrado) e a geometria da pagina.
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)
    return doc


def build():
    # carrega dados reais de Amapora e monta o mapa de substituicao
    global SUBST
    SUBST = build_subst_map()

    # Parte do template da 2S (timbrado correto), nao de um Document() vazio.
    doc = open_template_clean()
    style_base(doc)

    # --- secao 0: capa ---
    # A geometria de pagina e o timbrado vem do template (sectPr preservado).
    # Nao redefinimos margens/header_distance aqui para nao brigar com o
    # timbrado flutuante. Apenas garantimos o link do timbrado.
    build_cover(doc)

    # ====================================================================
    # CONTRACAPA / FOLHA DE IDENTIFICACAO
    # ====================================================================
    new_section(doc)
    h1(doc, "FOLHA DE IDENTIFICAÇÃO")
    add_para(doc,
             "Reúne-se nesta folha o conjunto de dados que identificam o empreendimento, este "
             "documento e os profissionais responsáveis por sua elaboração, verificação e "
             "aprovação, de forma a permitir o pronto reconhecimento do contexto contratual e "
             "técnico do Memorial Descritivo.", space_after=10)

    ident = [
        ("Empreendimento", "{{EMPREENDIMENTO}}"),
        ("Contratante", "ACCIONA"),
        ("Concessionária", "Companhia de Saneamento do Paraná – SANEPAR"),
        ("Projetista", "2S Engenharia e Geotecnologia"),
        ("Consórcio", "Consórcio E-Água e 2S Engenharia e Geotecnologia"),
        ("Objeto", "Projeto Executivo de Engenharia da Rede Coletora de Esgoto (RCE)"),
        ("Município", "{{MUNICIPIO}} – {{UF}}"),
        ("Sub-Bacia / Subsistema", "{{SUBBACIA}}"),
        ("Termo de Referência", "TR-PPP-SANEPAR-ACCIONA-07-2025-ENG"),
        ("ART nº", "{{ART}}"),
        ("Responsável Técnico", "{{ENG_RESP}}  –  CREA {{CREA}}"),
        ("Elaborado por", "{{ELABORADO}}"),
        ("Verificado por", "{{VERIFICADO}}"),
        ("Aprovado por", "{{APROVADO}}"),
        ("Código do Documento", "{{CODIGO_DOC}}"),
        ("Revisão", "{{REV}}"),
        ("Data", "{{DATA}}"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, (k, v) in enumerate(ident):
        cells = table.add_row().cells
        set_cell_bg(cells[0], "303030")
        set_cell_margins(cells[0]); set_cell_margins(cells[1])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pk = cells[0].paragraphs[0]; pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k); rk.bold = True; rk.font.size = Pt(10)
        rk.font.name = FONTE; rk.font.color.rgb = COR_BRANCO
        pv = cells[1].paragraphs[0]; pv.paragraph_format.space_after = Pt(0)
        if i % 2 == 1:
            set_cell_bg(cells[1], COR_CINZA_ZEBRA)
        rv = pv.add_run(v); rv.font.size = Pt(10); rv.font.name = FONTE
        cells[0].width = Cm(5.5); cells[1].width = Cm(10.0)

    # ====================================================================
    # CONTROLE DE REVISOES
    # ====================================================================
    new_section(doc)
    h1(doc, "CONTROLE DE REVISÕES")
    add_para(doc,
             "A tabela abaixo rastreia a evolução deste documento ao longo de suas emissões. "
             "Cada linha corresponde a uma revisão e indica, de forma sucinta, o que foi alterado, "
             "a data em que a versão foi emitida e os profissionais que a elaboraram e a aprovaram, "
             "garantindo a rastreabilidade exigida em projetos executivos.", space_after=10)
    add_table_caption(doc, "Controle de revisões do documento")
    make_table(
        doc,
        ["Rev.", "Descrição da Revisão", "Data", "Elaborou", "Aprovou"],
        [
            ["R0", "Emissão inicial", "{{DATA_R0}}", "{{ELABORADO}}", "{{APROVADO}}"],
            ["{{REV}}", "{{DESCRICAO_REV}}", "{{DATA}}", "{{ELABORADO}}", "{{APROVADO}}"],
            ["", "", "", "", ""],
        ],
        col_widths=[1.4, 6.6, 2.4, 3.0, 3.0],
        first_col_left=False,
    )
    add_table_source(doc)

    # ====================================================================
    # SUMARIO (TOC)
    # ====================================================================
    new_section(doc)
    h1(doc, "SUMÁRIO")
    p = doc.add_paragraph()
    add_toc(p, 'TOC \\o "1-3" \\h \\z \\u')

    # ====================================================================
    # LISTA DE FIGURAS / LISTA DE TABELAS  (montadas LITERALMENTE)
    # As listas sao preenchidas pelo gerador, com numeros literais; nao usam
    # campo TOC (que sairia vazio sem F9). Aqui apenas marcamos os pontos de
    # ancoragem; o conteudo e inserido em _fill_listas() ao final, quando
    # FIGURAS/TABELAS ja estao completas.
    # ====================================================================
    new_section(doc)
    h1(doc, "LISTA DE FIGURAS")
    anchor_fig = doc.add_paragraph()  # ancora: a lista de figuras entra apos esta
    h1(doc, "LISTA DE TABELAS")
    anchor_tab = doc.add_paragraph()  # ancora: a lista de tabelas entra apos esta

    # ====================================================================
    # 1. APRESENTACAO
    # ====================================================================
    new_section(doc)
    h1(doc, "1. APRESENTAÇÃO")
    add_para(doc,
             "Por meio de Parceria Público-Privada (PPP) celebrada com a Companhia de Saneamento do "
             "Paraná (SANEPAR), com prazo de 24 (vinte e quatro) anos, a ACCIONA passou a responder "
             "pelos serviços de esgotamento sanitário de 48 (quarenta e oito) municípios situados "
             "nas microrregiões Oeste e Centro-Leste do Estado do Paraná. O propósito dessa parceria "
             "é dar cumprimento às metas instituídas pelo Marco Legal do Saneamento Básico (Lei nº "
             "14.026/2020), avançando na universalização da coleta e do tratamento de esgoto em "
             "áreas hoje desprovidas de atendimento e, simultaneamente, ampliando e modernizando a "
             "infraestrutura já em operação.")
    add_para(doc,
             "Para viabilizar tais objetivos, firmou-se em 2025 um instrumento entre a "
             "concessionária ACCIONA e o Consórcio E-Água e 2S Engenharia e Geotecnologia, em "
             "atendimento ao Termo de Referência TR-PPP-SANEPAR-ACCIONA-07-2025-ENG. O ajuste tem "
             "por objeto a elaboração de projetos executivos de engenharia de sistemas de "
             "esgotamento sanitário nos municípios que compõem a Microrregião Oeste do Lote 2 da "
             "concessão SANEPAR, área sob responsabilidade da ACCIONA.")
    add_para(doc,
             "Inserido nesse contexto, este documento corresponde ao Memorial Descritivo do Projeto "
             "Executivo de Engenharia das redes coletoras de esgoto do subsistema {{SUBBACIA}}, parte "
             "integrante do Sistema de Esgotamento Sanitário (SES) do município de {{MUNICIPIO}} – "
             "{{UF}}. A abrangência do projeto contempla a sede municipal e os distritos "
             "{{DISTRITOS}}, segundo a delimitação das áreas de cobertura homologadas pela "
             "concessionária.")
    add_para(doc,
             "Ao longo dos capítulos seguintes são expostos o enquadramento do empreendimento, os "
             "critérios e parâmetros adotados no dimensionamento, as soluções construtivas, a "
             "metodologia de cálculo hidráulico e o resumo dos quantitativos obtidos. Todo o "
             "conteúdo observa as normas técnicas em vigor e as diretrizes de projeto da SANEPAR.")

    add_table_caption(doc, "Caracterização territorial e populacional – {{MUNICIPIO}} e área da {{SUBBACIA}}")
    make_table(
        doc,
        ["Caracterização", "Área (km²)", "População (hab)", "% do município"],
        [
            ["Município de {{MUNICIPIO}} – {{UF}}", "{{AREA_MUNI}}", "{{POP_MUNI}}", "100,00%"],
            ["Área da {{SUBBACIA}}", "{{AREA_SEDE}}", "{{POP_BACIA}}", "{{PCT_SEDE}}"],
        ],
        col_widths=[7.0, 3.5, 3.5, 3.0],
        first_col_left=True,
    )
    add_table_source(doc,
                     "Fonte: população e área territorial – IBGE, Censo 2022 (Amaporã-PR, "
                     "cód. 4100905); área da sub-bacia – envoltória da rede projetada (2S "
                     "Engenharia, 2026); população na sub-bacia estimada por nº de imóveis "
                     "atendidos × tamanho médio do domicílio.")
    add_para(doc,
             "A população residente do município de {{MUNICIPIO}} é de {{POP_MUNI}} habitantes "
             "(IBGE, Censo 2022) e sua área territorial é de {{AREA_MUNI}} km². A área da "
             "{{SUBBACIA}}, correspondente à envoltória da rede coletora projetada, é de "
             "{{AREA_SEDE}} km² ({{PCT_SEDE}} do território municipal), na qual se estima uma "
             "população atendida da ordem de {{POP_BACIA}} habitantes, obtida a partir dos "
             "imóveis atendidos pela rede.", space_before=6)

    # ====================================================================
    # 2. OBJETIVO E ESCOPO
    # ====================================================================
    new_section(doc)
    h1(doc, "2. OBJETIVO E ESCOPO")
    h2(doc, "2.1. Objetivo")
    add_para(doc,
             "Este Memorial Descritivo destina-se a apresentar, com o detalhamento técnico "
             "pertinente a um projeto executivo, a concepção e o dimensionamento da rede coletora de "
             "esgoto sanitário prevista para a {{SUBBACIA}}, no município de {{MUNICIPIO}} – {{UF}}. "
             "Para tanto, são reunidos os fundamentos normativos, os parâmetros de cálculo, as "
             "soluções construtivas e os quantitativos consolidados, fornecendo a base técnica "
             "necessária tanto à execução das obras quanto à sua análise e aprovação pela SANEPAR e "
             "pela concessionária ACCIONA.")
    h2(doc, "2.2. Escopo")
    add_para(doc, "Integram o escopo deste documento os seguintes temas:", space_after=4)
    add_bullet(doc, "a localização e a caracterização do sistema de esgotamento e da sub-bacia que para ele contribui;")
    add_bullet(doc, "o detalhamento da base topográfica e do referencial geodésico empregados no projeto;")
    add_bullet(doc, "os critérios e parâmetros adotados (vazões, coeficientes, material e diâmetro da tubulação, declividade e lâmina d'água);")
    add_bullet(doc, "as soluções construtivas previstas (traçado, recobrimento, tubos de queda, mudanças de diâmetro e acessórios);")
    add_bullet(doc, "a metodologia de dimensionamento hidráulico e o programa computacional utilizado;")
    if TEM_SOLEIRAS and TEM_INTERF:
        add_bullet(doc, "o tratamento das soleiras negativas e a verificação de interferências com redes existentes;")
    elif TEM_SOLEIRAS:
        add_bullet(doc, "o tratamento das soleiras negativas;")
    elif TEM_INTERF:
        add_bullet(doc, "a verificação de interferências com redes existentes;")
    add_bullet(doc, "a consolidação dos quantitativos da rede projetada, por diâmetro, por dispositivo e por faixa de profundidade.")
    add_para(doc,
             "O cálculo hidráulico trecho a trecho não é reproduzido neste texto: ele consta de "
             "documento próprio, a Planilha de Cálculo (Memorial de Cálculo – MC), que complementa "
             "o presente Memorial Descritivo.", space_before=6)

    # ====================================================================
    # 3. NORMAS E DOCUMENTOS DE REFERENCIA
    # ====================================================================
    new_section(doc)
    h1(doc, "3. NORMAS E DOCUMENTOS DE REFERÊNCIA")
    add_para(doc,
             "O desenvolvimento deste projeto teve por base o arcabouço normativo e os manuais "
             "setoriais listados a seguir, todos pertinentes ao projeto de sistemas de esgotamento "
             "sanitário por gravidade.", space_after=6)
    h2(doc, "3.1. Normas Técnicas (ABNT)")
    add_bullet(doc, "ABNT NBR 9649/1986 – Projeto de redes coletoras de esgoto sanitário – Procedimento;")
    add_bullet(doc, "ABNT NBR 14486/2000 – Sistemas enterrados para condução de esgoto sanitário – Projeto de redes coletoras com tubos de PVC;")
    add_bullet(doc, "ABNT NBR 12207 – Projeto de interceptores de esgoto sanitário;")
    add_bullet(doc, "ABNT NBR 9814 – Execução de rede coletora de esgoto sanitário.")
    h2(doc, "3.2. Manuais e Diretrizes da SANEPAR")
    add_bullet(doc, "MPS – Manual de Projetos de Saneamento (esgoto e água), SANEPAR;")
    add_bullet(doc, "MOS – Manual de Obras de Saneamento, SANEPAR;")
    add_bullet(doc, "Diretrizes para Elaboração de Projetos de Sistemas de Esgotamento Sanitário – Simulação Hidráulica, SANEPAR;")
    add_bullet(doc, "Nota Técnica de Tubulações para Redes de Sistemas de Abastecimento de Água (SAA) e de Esgotamento Sanitário (SES) – Requisitos, SANEPAR.")

    # ====================================================================
    # 4. CARACTERIZACAO E LOCALIZACAO
    # ====================================================================
    new_section(doc)
    h1(doc, "4. CARACTERIZAÇÃO E LOCALIZAÇÃO")
    h2(doc, "4.1. Caracterização do Sistema de Esgotamento Sanitário")
    add_para(doc,
             "Segundo o Estudo Técnico e Operacional do município de {{MUNICIPIO}}, o Sistema de "
             "Esgotamento Sanitário (SES) é estruturado a partir de redes coletoras, coletores "
             "principais e de uma Estação de Tratamento de Esgotos (ETE) denominada {{ETE}}. Com "
             "vistas a universalizar o atendimento, o referido estudo contempla a expansão das redes "
             "coletoras e a implantação de estações elevatórias de esgoto, com seus respectivos "
             "coletores de recalque, dimensionadas em função do relevo e da distribuição espacial das "
             "áreas de contribuição.")
    add_para(doc,
             "O subsistema aqui projetado, identificado como {{SUBBACIA}}, faz parte da sede "
             "municipal e encaminha suas contribuições por gravidade ao sistema coletor existente ou "
             "projetado, até o ponto de destinação para tratamento. Os limites da sub-bacia "
             "contribuinte foram definidos a partir do relevo local, do cadastro de imóveis "
             "levantado e das orientações da concessionária.")
    h2(doc, "4.2. Localização")
    add_para(doc,
             "Situado na microrregião Oeste do Estado do Paraná, o município de {{MUNICIPIO}} compõe "
             "o Lote 2 da concessão SANEPAR atribuído à ACCIONA. As figuras a seguir situam o "
             "município e a área de estudo no contexto regional e detalham os limites da "
             "{{SUBBACIA}} e o traçado da rede coletora projetada.")
    add_figure(doc, MAPA1, "Localização do município de {{MUNICIPIO}} – {{UF}} e da área de estudo")
    add_figure(doc, MAPA2, "Delimitação da {{SUBBACIA}} e traçado da rede coletora projetada")

    # ====================================================================
    # FLUXOGRAMA DO SISTEMA (pagina A4 PAISAGEM dedicada)
    # ====================================================================
    # TOLERANCIA A FLUXOGRAMA FALTANTE: so cria a pagina paisagem dedicada se o
    # PNG do fluxograma existir; senao, pula a secao inteira graciosamente (o
    # contador de figuras nao avanca, mantendo a numeracao correta).
    if FLUXOGRAMA and os.path.exists(FLUXOGRAMA):
        new_section_landscape(doc)
        _hflux = h1(doc, "FLUXOGRAMA DO SISTEMA")
        _hflux.paragraph_format.space_before = Pt(2)
        _hflux.paragraph_format.space_after = Pt(4)
        add_para(doc,
                 "O fluxograma a seguir sintetiza a concepção do Sistema de Esgotamento Sanitário da "
                 "{{SUBBACIA}}, desde a contribuição das redes coletoras, passando pela elevatória "
                 "({{EEE}}) e pela estação de tratamento ({{ETE}}), até a destinação final no corpo "
                 "receptor ({{CORPO_RECEPTOR}}).", space_after=4)
        # imagem grande ocupando a pagina landscape
        _fig_n[0] += 1
        _flux_n = _fig_n[0]
        FIGURAS.append((_flux_n, "Fluxograma do Sistema de Esgotamento Sanitário"))
        _sec = doc.sections[-1]
        _avail_w = page_width_emu(_sec)
        _avail_h = _sec.page_height - _sec.top_margin - _sec.bottom_margin
        _pflux = doc.add_paragraph()
        _pflux.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _pflux.paragraph_format.space_before = Pt(4)
        _pflux.paragraph_format.space_after = Pt(2)
        _rflux = _pflux.add_run()
        _picf = _rflux.add_picture(FLUXOGRAMA, width=Emu(int(_avail_w * 0.90)))
        # cap de altura conservador para que TODO o conteudo da pagina paisagem
        # (titulo + intro + imagem + legenda + fonte + paragrafo de quebra de
        # secao) caiba numa UNICA pagina, evitando a pagina paisagem em branco.
        _maxhf = int(_avail_h * 0.66)
        if _picf.height > _maxhf:
            _rt = _maxhf / _picf.height
            _picf.height = Emu(int(_picf.height * _rt))
            _picf.width = Emu(int(_picf.width * _rt))
        # legenda como estilo Caption + campo SEQ Figura (popula a Lista de Figuras)
        add_caption(doc, "Figura", _flux_n, "Fluxograma do Sistema de Esgotamento Sanitário",
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        _srcf = doc.add_paragraph()
        _srcf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rsf = _srcf.add_run("Fonte: 2S Engenharia, 2026.")
        _rsf.italic = True; _rsf.font.size = Pt(8.5); _rsf.font.name = FONTE
        _rsf.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    else:
        sys.stderr.write(
            "[memorial] AVISO: fluxograma ausente, secao pulada -> %s\n" % FLUXOGRAMA)

    # ====================================================================
    # 5. BASE TOPOGRAFICA E GEORREFERENCIAMENTO  (volta a RETRATO)
    # ====================================================================
    new_section_portrait(doc)
    h1(doc, "5. BASE TOPOGRÁFICA E GEORREFERENCIAMENTO")
    add_para(doc,
             "O lançamento e o dimensionamento da rede coletora apoiaram-se no levantamento "
             "topográfico cadastral georreferenciado realizado pela 2S Engenharia e Geotecnologia, "
             "integrante do consórcio projetista. Esse levantamento abrangeu o cadastro das vias, "
             "dos imóveis (com as respectivas cotas de soleira e de edificação), das interferências "
             "aparentes e das redes de infraestrutura existentes informadas pelos órgãos "
             "competentes.")
    # Sintese numerica do levantamento: so com dados reais de TXT (TEM_TOPO).
    if TEM_TOPO:
        add_para(doc,
                 "O levantamento da área da {{SUBBACIA}} totalizou {{TOPO_N_PONTOS}} pontos topográficos "
                 "cadastrais, com cotas de terreno variando entre {{TOPO_Z_MIN}} m e {{TOPO_Z_MAX}} m (cota média de "
                 "{{TOPO_Z_MED}} m), o que resulta em um desnível total da ordem de {{TOPO_DESNIVEL}} m na área de estudo. A "
                 "figura a seguir apresenta o mapa hipsométrico (mapa de calor da topografia) elaborado "
                 "a partir desses pontos, evidenciando a conformação do relevo que condicionou o traçado "
                 "e as declividades da rede coletora.")
        add_figure(doc, MAPA4, "Mapa hipsométrico da {{SUBBACIA}} (modelo de calor da topografia)")
        add_figure(doc, MAPA6,
                   "Modelo digital do terreno em vista isométrica (exagero vertical 4×)")
        add_table_caption(doc, "Síntese do levantamento topográfico da {{SUBBACIA}}")
        make_table(
            doc,
            ["Parâmetro", "Valor"],
            [
                ["Pontos topográficos levantados", "{{TOPO_N_PONTOS}}"],
                ["Cota mínima do terreno", "{{TOPO_Z_MIN}} m"],
                ["Cota máxima do terreno", "{{TOPO_Z_MAX}} m"],
                ["Cota média do terreno", "{{TOPO_Z_MED}} m"],
                ["Desnível total", "{{TOPO_DESNIVEL}} m"],
                ["Datum / Sistema de projeção", "SIRGAS 2000 / UTM 22S (EPSG:31982)"],
                ["Executor do levantamento", "2S Engenharia e Geotecnologia"],
            ],
            col_widths=[8.0, 7.0],
            first_col_left=True,
        )
        add_table_source(doc)
    h2(doc, "5.1. Sistema de Referência")
    add_bullet(doc, "Sistema Geodésico de Referência: SIRGAS 2000;")
    add_bullet(doc, "Sistema de projeção: UTM – Universal Transversa de Mercator;")
    add_bullet(doc, "Fuso / Meridiano Central: Fuso 22 Sul (MC 51º W);")
    add_bullet(doc, "Referencial altimétrico: {{REF_ALTIMETRICO}} (Datum vertical – Imbituba/SC ou marco oficial local);")
    add_bullet(doc, "Unidade: metros (m).")
    h2(doc, "5.2. Métodos e Equipamentos")
    add_para(doc,
             "O cadastro foi executado por posicionamento GNSS pelo método {{GNSS_METODO}}, "
             "referenciado a bases de coordenadas conhecidas, complementado, quando necessário, por "
             "topografia convencional com estação total para o adensamento e o cadastro de detalhes. "
             "O levantamento totalizou {{GNSS_N_PONTOS}} pontos cadastrais, com PDOP médio da ordem "
             "de {{GNSS_PDOP}}, condição geométrica favorável à obtenção de soluções fixas de alta "
             "qualidade. A técnica efetivamente adotada em cada trecho consta do Relatório Técnico "
             "de Levantamento Topográfico.")
    h2(doc, "5.3. Precisões Obtidas")
    add_bullet(doc, "Precisão planimétrica típica: {{PREC_PLANI}};")
    add_bullet(doc, "Precisão altimétrica típica: {{PREC_ALTI}};")
    add_bullet(doc, "Método de posicionamento: {{GNSS_METODO}};")
    add_bullet(doc, "Número de pontos levantados: {{GNSS_N_PONTOS}};")
    add_bullet(doc, "Pontos de apoio / RRNN implantados: {{QTD_APOIO}}.")
    add_para(doc,
             "Os valores acima correspondem à precisão típica (mediana) do conjunto de pontos. "
             "Pontos situados sob copa de árvores, junto a edificações ou em demais áreas de "
             "obstrução parcial do sinal dos satélites apresentam precisão inferior à típica, "
             "conforme usual em levantamentos GNSS em meio urbano arborizado.", space_before=4)
    add_para(doc,
             "Os níveis de precisão alcançados satisfazem as exigências de projetos executivos de "
             "redes coletoras, conferindo confiabilidade às cotas de terreno e aos perfis "
             "longitudinais que serviram de base ao dimensionamento.", space_before=4)

    # ====================================================================
    # 6. INVESTIGACAO GEOTECNICA (SONDAGENS)
    # ====================================================================
    if SOND:
        _sp = bool(SOND.get("tem_spt"))
        _na = bool(SOND.get("tem_na"))
        new_section_portrait(doc)
        h1(doc, "6. INVESTIGAÇÃO GEOTÉCNICA (SONDAGENS)")
        add_para(doc,
                 "Como parte do desenvolvimento do projeto da rede coletora de esgoto (RCE) do "
                 "município de {{MUNICIPIO}}, foi conduzida campanha de investigação geotécnica do "
                 "subsolo ao longo do traçado projetado, com a finalidade de caracterizar os horizontes "
                 "de solo interceptados pelas obras e subsidiar as definições de projeto e de execução. "
                 "A investigação totalizou {{SOND_N_FUROS}} pontos de reconhecimento, executados por "
                 "{{SOND_METODO_INVEST}}, com profundidade investigada de {{SOND_PROF_MIN}} m a "
                 "{{SOND_PROF_MAX}} m (média de {{SOND_PROF_MED}} m), distribuídos de forma a representar "
                 "as condições do subsolo nos trechos de maior interesse executivo.")
        add_para(doc,
                 "Os resultados destinam-se a subsidiar, em especial: a definição do método executivo de "
                 "abertura das valas; a avaliação da estabilidade das paredes de escavação e a "
                 "necessidade de escoramento e/ou execução em talude; a verificação da necessidade de "
                 "esgotamento ou rebaixamento do lençol freático; o assentamento da tubulação sobre berço "
                 "e envoltória adequados; o dimensionamento das condições de apoio e fundação dos poços "
                 "de visita (PV); e a avaliação preliminar da agressividade do solo aos materiais "
                 "empregados.")
        add_para(doc,
                 "A campanha observou as prescrições normativas pertinentes, com destaque para a ABNT "
                 "NBR 8036 (programação de sondagens), a NBR 9603 (sondagem a trado), a NBR 6484 "
                 "(sondagem de simples reconhecimento com SPT) e a NBR 6502 (terminologia de solos). A "
                 "classificação e a descrição das amostras seguiram a terminologia consagrada nessas "
                 "normas.")
        add_figure(doc, MAPA7, "Mapa de locação das sondagens — {{MUNICIPIO}} – {{UF}}")
        add_table_caption(doc, "Síntese da investigação geotécnica — {{MUNICIPIO}}")
        make_table(
            doc, ["Parâmetro", "Valor"],
            [
                ["Número de furos / pontos investigados", "{{SOND_N_FUROS}}"],
                ["Método(s) de investigação", "{{SOND_METODO_INVEST}}"],
                ["Profundidade investigada – mínima", "{{SOND_PROF_MIN}} m"],
                ["Profundidade investigada – média", "{{SOND_PROF_MED}} m"],
                ["Profundidade investigada – máxima", "{{SOND_PROF_MAX}} m"],
                ["Solo predominante", "{{SOND_SOLO_PRED}}"],
                ["Cor predominante", "{{SOND_COR_PRED}}"],
                ["Ensaio SPT realizado", "{{SOND_SPT_SIM_NAO}}"],
                ["Nível d'água", "{{SOND_NA_SITUACAO}}"],
                ["Datum da locação (GNSS)", "{{SOND_DATUM}}"],
            ],
            col_widths=[8.0, 7.0], first_col_left=True,
        )
        add_table_source(doc)

        h2(doc, "6.1. Metodologia e Abrangência")
        add_para(doc,
                 "A investigação compreendeu {{SOND_N_FUROS}} pontos de reconhecimento do subsolo, "
                 "executados por {{SOND_METODO_INVEST}}, ao longo do traçado da rede coletora projetada. "
                 "A locação de cada ponto foi amarrada por posicionamento por satélite (GNSS), no datum "
                 "{{SOND_DATUM}}, garantindo a correlação espacial entre os pontos investigados e as "
                 "estruturas projetadas (coletores, PVs e órgãos acessórios), conforme o mapa de locação "
                 "apresentado acima.")
        if _sp:
            add_para(doc,
                     "Parte dos pontos foi investigada por sondagem de simples reconhecimento com ensaio "
                     "de penetração dinâmica (SPT), em conformidade com a NBR 6484, obtendo-se o índice "
                     "de resistência à penetração (NSPT) ao longo da profundidade; os demais foram "
                     "reconhecidos por sondagem a trado e/ou poços de inspeção (NBR 9603).", space_before=4)
        else:
            add_para(doc,
                     "A investigação foi conduzida por sondagem a trado e/ou abertura de poços de "
                     "inspeção, em conformidade com a NBR 9603 — método adequado à profundidade de "
                     "interesse das obras de rede coletora, predominantemente rasas. Em razão do método, "
                     "não foram obtidos índices de resistência à penetração (SPT); a caracterização da "
                     "compacidade/consistência baseou-se na identificação táctil-visual e no comportamento "
                     "do solo durante a perfuração.", space_before=4)

        h2(doc, "6.2. Perfil Estratigráfico do Subsolo")
        add_para(doc,
                 "O reconhecimento do subsolo indicou predominância de solo {{SOND_SOLO_L}} de coloração "
                 "{{SOND_COR_L}}, perfil característico da região onde se insere o município de "
                 "{{MUNICIPIO}}. As camadas investigadas apresentaram, em geral, homogeneidade ao longo "
                 "da profundidade reconhecida, sem variações estratigráficas significativas até a "
                 "profundidade investigada.")
        add_para(doc,
                 "Do ponto de vista da geologia regional, a área situa-se no domínio do Grupo Bauru / "
                 "Formação Caiuá (arenitos da Bacia do Paraná), sobre o qual se desenvolvem, no noroeste "
                 "do Paraná, solos areno-argilosos de natureza laterítica (latossolos e argissolos "
                 "vermelhos), resultantes do intemperismo do substrato arenítico. Tais solos "
                 "caracterizam-se por matriz predominantemente arenosa fina, coloração avermelhada "
                 "conferida por óxidos e hidróxidos de ferro, boa drenagem natural e baixa plasticidade — "
                 "caracterização coerente com o perfil verificado em campo.", space_before=4)

        h2(doc, "6.3. Compacidade e Consistência")
        if _sp:
            add_para(doc,
                     "Os pontos investigados com SPT (NBR 6484) tiveram o estado de compacidade (solos "
                     "arenosos) e de consistência (solos argilosos) classificados em função do índice "
                     "NSPT, conforme os critérios da NBR 6484/7250, reproduzidos a seguir.")
            add_table_caption(doc, "Compacidade de solos arenosos em função do NSPT (NBR 6484)")
            make_table(
                doc, ["NSPT (golpes)", "Designação"],
                [["≤ 4", "Fofa"], ["5 a 8", "Pouco compacta"], ["9 a 18", "Medianamente compacta"],
                 ["19 a 40", "Compacta"], ["> 40", "Muito compacta"]],
                col_widths=[7.0, 8.0], first_col_left=True,
            )
            add_table_caption(doc, "Consistência de solos argilosos em função do NSPT (NBR 6484)")
            make_table(
                doc, ["NSPT (golpes)", "Designação"],
                [["≤ 2", "Muito mole"], ["3 a 5", "Mole"], ["6 a 10", "Média"],
                 ["11 a 19", "Rija"], ["> 19", "Dura"]],
                col_widths=[7.0, 8.0], first_col_left=True,
            )
            add_para(doc,
                     "Em síntese, os resultados de SPT enquadram o solo investigado predominantemente "
                     "como {{SOND_COMPACIDADE_PRED}} (NSPT médio da ordem de {{SOND_NSPT_MED}}), "
                     "conforme os boletins individuais de sondagem.", space_before=4)
        else:
            add_para(doc,
                     "Em razão de a investigação ter sido conduzida por sondagem a trado / poços de "
                     "inspeção (NBR 9603), não há determinação do índice de resistência à penetração "
                     "(SPT) para este município. A estimativa do estado de compacidade dos solos arenosos "
                     "e de consistência dos solos argilosos apoiou-se na identificação táctil-visual e na "
                     "resistência oferecida à perfuração, classificadas conforme a terminologia da "
                     "NBR 6502.")

        h2(doc, "6.4. Nível d'Água")
        if _na:
            add_para(doc,
                     "A investigação detectou a presença de nível d'água em {{SOND_N_FUROS_NA}} ponto(s), "
                     "à profundidade média de {{SOND_NA_PROF_MED}} m em relação ao terreno. A presença de "
                     "lençol freático na faixa de escavação impõe atenção ao método executivo, demandando "
                     "previsão de esgotamento e/ou rebaixamento nos trechos afetados, bem como cuidados "
                     "quanto à estabilidade das paredes de vala e ao carreamento de finos. As cotas do "
                     "nível d'água por ponto constam dos boletins anexos.")
        else:
            add_para(doc,
                     "Durante a execução da investigação não foi detectado nível d'água até a "
                     "profundidade investigada ({{SOND_PROF_MAX}} m). A ausência de lençol freático na "
                     "faixa de interesse das obras, associada à natureza arenosa e bem drenada do solo, "
                     "configura condição favorável à execução das valas, dispensando, em princípio, "
                     "medidas de rebaixamento ou esgotamento. Ressalva-se que o nível d'água é variável "
                     "sazonalmente, podendo elevar-se em períodos chuvosos, recomendando-se a observação "
                     "das condições de campo por ocasião da execução.")

        h2(doc, "6.5. Implicações Geotécnicas para a Obra")
        add_para(doc,
                 "Estabilidade das valas. O solo predominante {{SOND_SOLO_L}}, embora favorável quanto à "
                 "escavabilidade, apresenta baixa coesão e paredes de escavação instáveis quando deixadas "
                 "na vertical. Para valas com profundidade superior a aproximadamente 1,25 m é mandatória "
                 "a adoção de escoramento das paredes ou execução em talude compatível com o ângulo de "
                 "atrito do material, em atendimento à NR-18 e à NBR 9061, com atenção redobrada em "
                 "escavações próximas a edificações, vias e interferências.")
        if _na:
            add_para(doc,
                     "Presença de água. Verificada a presença de nível d'água, os trechos afetados "
                     "demandarão esgotamento e/ou rebaixamento prévio, além de cuidados contra o "
                     "carreamento de finos e a instabilização do fundo de vala, agravados em solos "
                     "arenosos saturados.", space_before=4)
        else:
            add_para(doc,
                     "Ausência de água. Não tendo sido detectado nível d'água até a profundidade "
                     "investigada, prevê-se escavação a seco, sem necessidade, a princípio, de "
                     "rebaixamento, mantida a recomendação de monitoramento sazonal.", space_before=4)
        add_para(doc,
                 "Berço, envoltória e reaterro. O assentamento da tubulação deverá apoiar-se sobre berço "
                 "e envoltória de material granular selecionado, conforme padrão SANEPAR, com reaterro "
                 "executado em camadas compactadas, isento de matéria orgânica, em especial sob vias e "
                 "travessias. O solo arenoso local, bem graduado e compactado, presta-se adequadamente a "
                 "reaterro, controlada a umidade de compactação.", space_before=4)
        add_para(doc,
                 "Fundação dos poços de visita e agressividade. A fundação dos PVs deverá apoiar-se em "
                 "terreno de capacidade de suporte compatível, regularizando-se e compactando-se o fundo "
                 "de escavação. O solo {{SOND_SOLO_L}} de coloração {{SOND_COR_L}}, de natureza laterítica "
                 "e bem drenado, caracteriza-se por baixa agressividade aos materiais constituintes da "
                 "rede (tubulações e estruturas de concreto), não sendo esperada condição de ataque "
                 "químico relevante, observadas as especificações de materiais do empreendimento.",
                 space_before=4)

        h2(doc, "6.6. Conclusão Geotécnica")
        add_para(doc,
                 "A investigação realizada no município de {{MUNICIPIO}}, composta por {{SOND_N_FUROS}} "
                 "pontos executados por {{SOND_METODO_INVEST}} até {{SOND_PROF_MAX}} m, caracterizou o "
                 "subsolo como predominantemente {{SOND_SOLO_L}} de coloração {{SOND_COR_L}}, perfil "
                 "típico dos solos lateríticos do noroeste do Paraná associados à Formação Caiuá.")
        if _na:
            add_para(doc,
                     "A presença de nível d'água em parte dos pontos configura condição condicionada, "
                     "exigindo medidas de esgotamento/rebaixamento e cuidados executivos nos trechos "
                     "afetados.", space_before=4)
        else:
            add_para(doc,
                     "Não foi detectado nível d'água até a profundidade investigada, o que, somado à "
                     "natureza arenosa e bem drenada do material, configura condição geotécnica "
                     "globalmente favorável à execução da rede coletora.", space_before=4)
        add_bullet(doc, "Adoção obrigatória de escoramento ou taludamento das valas com profundidade superior a aproximadamente 1,25 m (NR-18);")
        add_bullet(doc, "Verificação das condições de fundo de escavação para apoio dos poços de visita;")
        add_bullet(doc, "Execução de berço/envoltória e reaterro compactado conforme padrão SANEPAR;")
        add_bullet(doc, "Observação das condições de campo, sobretudo do nível d'água, por ocasião da execução;")
        add_bullet(doc, "Complementação por sondagens com SPT nos trechos mais profundos ou de condição mais crítica, quando pertinente.")

    # ====================================================================
    # 7. CRITERIOS E PARAMETROS DE PROJETO
    # ====================================================================
    new_section(doc)
    h1(doc, "7. CRITÉRIOS E PARÂMETROS DE PROJETO")
    add_para(doc,
             "O cálculo das tubulações em regime de escoamento livre foi orientado pelas normas da "
             "ABNT (NBR 9649/1986 e NBR 14486/2000) e pelas diretrizes técnicas da SANEPAR, "
             "harmonizadas com as premissas pactuadas com a concessionária ACCIONA. Os itens a "
             "seguir descrevem, um a um, os critérios efetivamente adotados.")

    h2(doc, "7.1. Vazões Sanitárias de Contribuição")
    add_para(doc,
             "No início de operação, a vazão mínima de contribuição resulta da soma da vazão "
             "sanitária mínima com a parcela de infiltração, ambas referidas ao ano fixado como "
             "início do plano. Já a vazão máxima de contribuição corresponde àquela gerada pela "
             "população máxima prevista para o horizonte de projeto, nas condições de dia e de hora "
             "de maior consumo, somada à respectiva vazão de infiltração.")
    add_para(doc,
             "A vazão média de esgoto é calculada por Qmed = (C · P · q) / 86.400, na qual C "
             "representa o coeficiente de retorno, P a população contribuinte e q o consumo per "
             "capita, em L/hab.dia. A partir dela, a vazão máxima horária decorre de "
             "Qmax,h = K1 · K2 · Qmed e a vazão mínima horária de Qmin = K3 · Qmed, sendo K1, K2 e "
             "K3 os coeficientes que traduzem a variação temporal das vazões (dia de maior consumo, "
             "hora de maior consumo e hora de menor consumo, respectivamente).")
    add_para(doc,
             "Como salvaguarda normativa, sempre que a vazão calculada em um trecho resultar inferior "
             "a 1,50 L/s, o dimensionamento é refeito com a adoção desse valor de 1,50 L/s como "
             "vazão mínima de projeto.")

    h2(doc, "7.2. Vazões de Infiltração")
    add_para(doc,
             "De acordo com as ABNT NBR 9649 e NBR 14486, a taxa de infiltração é função de "
             "condicionantes locais — entre eles a posição do lençol freático, a natureza do "
             "subsolo, o esmero na execução da rede, o material da tubulação e o tipo de junta "
             "empregada. A correspondente vazão de infiltração resulta do produto dessa taxa pela "
             "extensão total da rede coletora prevista.")
    add_para(doc,
             "A ABNT NBR 9649/1986 admite, para a contribuição por infiltração, valores entre 0,05 e "
             "1,00 L/s·km, conforme as condições de cada local. Neste projeto, em alinhamento com a "
             "SANEPAR e com o estudo de concepção, adotou-se a taxa de {{TX_INFILTRACAO}} L/s·km, "
             "compatível com as características locais e com as premissas de planejamento do "
             "sistema.")

    h2(doc, "7.3. Per Capita e Coeficientes")
    add_para(doc, "Para o cálculo das vazões de projeto foram adotados os seguintes parâmetros:", space_after=4)
    add_table_caption(doc, "Per capita e coeficientes adotados no projeto")
    make_table(
        doc,
        ["Parâmetro", "Símbolo", "Valor", "Unidade"],
        [
            ["Consumo per capita", "q (QPERC)", "{{QPERC}}", "L/hab.dia"],
            ["Coeficiente do dia de maior consumo", "K1", "{{K1}}", "–"],
            ["Coeficiente da hora de maior consumo", "K2", "{{K2}}", "–"],
            ["Coeficiente da hora de menor consumo", "K3", "{{K3}}", "–"],
            ["Coeficiente de retorno", "C", "{{C_RETORNO}}", "–"],
            ["Taxa de infiltração", "Tinf", "{{TX_INFILTRACAO}}", "L/s·km"],
        ],
        col_widths=[7.0, 3.0, 3.0, 3.0],
        first_col_left=True,
    )
    add_table_source(doc)

    h2(doc, "7.4. Tubulação Adotada")
    add_para(doc,
             "A Nota Técnica de Tubulações da SANEPAR relaciona, como materiais admissíveis em "
             "sistemas de esgotamento por gravidade, o PVC, o PEAD, o PVC corrugado e o PP "
             "corrugado. Observado o que foi acordado com a concessionária ACCIONA, a rede coletora "
             "foi dimensionada adotando-se:")
    add_bullet(doc, "Material da tubulação: PVC – Policloreto de Vinila (PVC JEI – junta elástica integrada);")
    add_bullet(doc, "Diâmetro nominal mínimo: DN 150 mm;")
    add_bullet(doc, "Coeficiente de Manning: n = 0,010;")
    add_bullet(doc, "Tensão trativa mínima: 1,00 Pa.")
    add_para(doc,
             "Embora a ABNT NBR 9649 traga, como valor de referência corrente, o coeficiente de "
             "rugosidade de Manning n = 0,013, optou-se, para as tubulações em PVC, pelo valor "
             "n = 0,010. Essa escolha apoia-se na ABNT NBR 14486 — norma dedicada especificamente às "
             "redes coletoras executadas com tubos de PVC —, que reconhece o desempenho hidráulico "
             "superior desse material, em razão, sobretudo, de sua reduzida rugosidade interna.",
             space_before=6)
    add_para(doc,
             "O mesmo valor encontra amparo na literatura técnica de referência do setor, em "
             "particular na obra Coleta e Transporte de Esgoto Sanitário, de Milton Tomoyuki "
             "Tsutiya, que associa aos tubos de PVC o coeficiente de Manning igual a 0,010. Ademais, "
             "n = 0,010 coincide com os valores que as bibliotecas hidráulicas dos programas de "
             "simulação (SewerCAD/SewerGEMS) atribuem ao PVC, justamente em função de sua parede "
             "interna lisa e da baixa tendência à incrustação. Conclui-se, portanto, que a adoção "
             "desse coeficiente é tecnicamente consistente e aderente às normas específicas do "
             "material utilizado.")

    h2(doc, "7.5. Declividade")
    add_para(doc,
             "A declividade mínima de cada trecho foi estabelecida segundo a ABNT NBR 14486/2000, "
             "que condiciona a inclinação do conduto ao atendimento da tensão trativa mínima de "
             "autolimpeza. Em termos práticos, a inclinação mínima é fixada de modo a garantir essa "
             "tensão trativa, sendo comumente expressa, para a vazão inicial, pela relação "
             "Imin = 0,0055 · Qi^(-0,47), na qual Qi é a vazão inicial em L/s.")
    add_para(doc,
             "Tomando-se a vazão mínima de projeto de 1,50 L/s, a expressão conduz a uma declividade "
             "mínima da ordem de {{DECL_MIN}} m/m, suficiente para satisfazer as condições "
             "hidráulicas normativas. No outro extremo, a declividade máxima foi limitada de modo a "
             "manter a velocidade do escoamento até 5,00 m/s, evitando-se o desgaste prematuro das "
             "tubulações e a ocorrência de instabilidades hidráulicas.")
    add_para(doc,
             "Por fim, em consonância com as diretrizes pactuadas com a concessionária e buscando "
             "soluções de execução mais viáveis, o dimensionamento privilegiou, sempre que possível, "
             "a declividade mínima de 1,00% (0,01 m/m) associada à tensão trativa mínima de "
             "1,00 Pa.")

    h2(doc, "7.6. Altura da Lâmina d'Água")
    add_para(doc,
             "A ABNT NBR 9649/1986 fixa, para regime uniforme e permanente, a lâmina máxima de "
             "escoamento em 75% do diâmetro interno do coletor, avaliada na condição de vazão final "
             "de projeto. Esse foi o critério hidráulico adotado: limitar a lâmina d'água a, no "
             "máximo, 75% do diâmetro quando submetida à vazão final.")
    add_para(doc,
             "Como medida adicional de segurança, e seguindo orientação da concessionária, "
             "estabeleceu-se que, para o coletor principal — aquele que concentra as maiores vazões "
             "do sistema —, todo trecho cujo cálculo indique lâmina superior a 50% do diâmetro "
             "interno deve ter seu diâmetro nominal majorado. Tal providência aumenta a folga "
             "operacional, mitiga o risco de sobrecarga hidráulica e preserva o desempenho do "
             "coletor ao longo de todo o horizonte de projeto.")

    # ====================================================================
    # 7. DISPOSICOES CONSTRUTIVAS
    # ====================================================================
    new_section(doc)
    h1(doc, "8. DISPOSIÇÕES CONSTRUTIVAS")

    h2(doc, "8.1. Traçado da Rede Coletora")
    add_para(doc,
             "O traçado da rede coletora foi concebido de acordo com o Manual de Obras de Saneamento "
             "(MOS) da SANEPAR. Por se tratar de subsistema implantado em área com pavimentação "
             "definitiva, optou-se por arranjo duplo, com as redes locadas nos passeios, mantendo-se "
             "afastamento mínimo de 1,50 m em relação ao alinhamento predial e recobrimento mínimo "
             "de 0,95 m. Além desses, observaram-se os critérios listados a seguir:")
    add_bullet(doc, "Distância máxima entre dispositivos de inspeção limitada a 100 m, em atendimento aos critérios operacionais e de manutenção;")
    add_bullet(doc, "Profundidade máxima das redes coletoras limitada a 3,00 m, ressalvadas situações especiais devidamente justificadas e acordadas com a concessionária;")
    add_bullet(doc, "Lançamento da rede com base no levantamento topográfico cadastral executado pela 2S Engenharia e Geotecnologia.")

    h2(doc, "8.2. Recobrimento e Profundidades")
    add_para(doc,
             "Sobre a geratriz superior dos tubos assentados em passeio adotou-se recobrimento "
             "mínimo de 0,95 m, valor que protege a canalização das cargas externas e ainda "
             "viabiliza as ligações prediais por gravidade. A profundidade de projeto foi, em regra, "
             "limitada a 3,00 m; apenas em situações excepcionais e devidamente justificadas "
             "admitiu-se chegar a 3,50 m, conforme faculta o MOS/SANEPAR e mediante anuência da "
             "concessionária.")

    h2(doc, "8.3. Tubo de Queda e Degrau")
    add_para(doc,
             "Prevê-se a instalação de tubo de queda (TQ) toda vez que o desnível entre o fundo do "
             "poço de visita e a geratriz inferior da tubulação de montante ultrapassar 0,50 m. "
             "Quando viável, os degraus apontados pelo cálculo são suprimidos elevando-se a "
             "declividade do trecho de montante, de sorte que a geratriz inferior do tubo passe a "
             "coincidir com a cota de fundo do poço de visita.")
    add_detalhe(doc, DET_TQ, "Detalhamento de Tubo de Queda (TQ)")

    h2(doc, "8.4. Mudança de Diâmetro")
    add_para(doc,
             "As diretrizes da SANEPAR fixam o diâmetro nominal mínimo em DN 150. Até o DN 400 "
             "especifica-se o PVC com junta elástica integrada (PVC JEI); para diâmetros maiores ou "
             "regime pressurizado, procede-se a uma análise comparativa entre as alternativas "
             "disponíveis no mercado. Como a vazão cresce de montante para jusante, não se admite "
             "reduzir o diâmetro interno em trechos a jusante; eventuais reduções só podem ser "
             "feitas com justificativa técnica e prévia anuência da SANEPAR.")
    add_para(doc,
             "Havendo mudança de diâmetro, rebaixa-se a geratriz inferior interna do tubo de maior "
             "diâmetro, de modo a preservar a continuidade do gradiente hidráulico, conforme a "
             "expressão:")
    eq = doc.add_paragraph()
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    req = eq.add_run("r = 0,75 · (D – d)")
    req.bold = True
    req.font.size = Pt(12)
    req.font.name = FONTE
    eq.paragraph_format.space_after = Pt(4)
    add_para(doc, "onde: r = rebaixo da tubulação de maior diâmetro; D e d = diâmetros maior e menor, "
                  "respectivamente.", size=10, space_before=0)

    h2(doc, "8.5. Acessórios da Rede Coletora")
    h3(doc, "7.5.1. Poço de Visita (PV)")
    add_para(doc,
             "Os poços de visita são previstos em todos os pontos singulares da rede — início de "
             "coletores, mudanças de alinhamento horizontal ou vertical, alterações de declividade, "
             "de diâmetro ou de material, encontros de coletores e locais que demandem degraus "
             "hidráulicos ou tubos de queda. Sua adoção torna-se obrigatória, em especial:")
    add_bullet(doc, "na reunião de mais de dois trechos ao coletor;")
    add_bullet(doc, "em reuniões que exijam a implantação de tubo de queda;")
    add_bullet(doc, "nas extremidades de sifões invertidos e passagens forçadas;")
    add_bullet(doc, "no início de coletores, em substituição a terminais, quando a profundidade da tubulação for igual ou superior a 3,00 m.")
    add_para(doc,
             "Levando em conta o alcance útil dos equipamentos de limpeza e desobstrução "
             "(mangueiras de hidrojateamento), o espaçamento entre dois PV consecutivos foi "
             "mantido, sempre que tecnicamente exequível, em no máximo 80,00 m.", space_before=4)
    add_detalhe(doc, DET_PV_A, "Detalhamento de Poço de Visita Tipo A")
    add_detalhe(doc, DET_PV1000, "Detalhamento de Poço de Visita PV 1000")
    add_detalhe(doc, DET_PVTR, "Detalhamento de Poço de Visita com Tubo de Redução (PVTR)")
    h3(doc, "7.5.2. Terminal de Limpeza (TL)")
    add_para(doc,
             "Nos trechos iniciais de coletores, o terminal de limpeza pode fazer as vezes do poço "
             "de visita, oferecendo acesso dos equipamentos de limpeza ao trecho situado a jusante.")
    add_detalhe(doc, DET_TL, "Detalhamento de Terminal de Limpeza (TL)")
    h3(doc, "7.5.3. Terminal de Inspeção e Limpeza (TIL)")
    add_para(doc,
             "O terminal de inspeção e limpeza também pode substituir o poço de visita, desde que "
             "satisfeitas as condições técnicas e operacionais, nas seguintes situações:")
    add_bullet(doc, "para possibilitar o acesso de equipamentos de limpeza ao trecho a jusante da rede;")
    add_bullet(doc, "no início de coletores, quando não houver necessidade de poço de visita;")
    add_bullet(doc, "na reunião de até dois trechos em um mesmo coletor;")
    add_bullet(doc, "em pontos onde ocorra degrau hidráulico com altura inferior a 0,50 m;")
    add_bullet(doc, "a jusante de ligações prediais cujas contribuições possam ocasionar dificuldades operacionais.")

    h3(doc, "7.5.4. Ligações Domiciliares")
    add_para(doc,
             "A interligação de cada imóvel à rede coletora pública é feita por ligação domiciliar, "
             "executada conforme os padrões da SANEPAR e adaptada à posição relativa entre a soleira "
             "do imóvel e a rede. São previstos três tipos construtivos de ligação domiciliar, "
             "ilustrados a seguir, selecionados em função da profundidade da rede e da cota da "
             "edificação.")
    add_detalhe(doc, DET_LIG1, "Detalhamento de Ligação Domiciliar – Tipo 01")
    add_detalhe(doc, DET_LIG2, "Detalhamento de Ligação Domiciliar – Tipo 02")
    add_detalhe(doc, DET_LIG3, "Detalhamento de Ligação Domiciliar – Tipo 03")

    # ====================================================================
    # 8. DIMENSIONAMENTO HIDRAULICO
    # ====================================================================
    new_section(doc)
    h1(doc, "9. DIMENSIONAMENTO HIDRÁULICO")

    h2(doc, "9.1. Metodologia e Software")
    add_para(doc,
             "O dimensionamento da rede coletora seguiu o Manual de Obras de Saneamento (MOS) da "
             "SANEPAR, as normas ABNT NBR 9649/1986 e NBR 14486/2000 e as diretrizes validadas com a "
             "ACCIONA. Em atendimento às Diretrizes para Elaboração de Projetos de Sistemas de "
             "Esgotamento Sanitário – Simulação Hidráulica, editadas pela SANEPAR, a modelagem foi "
             "executada em programa específico, integralmente compatível com a plataforma da "
             "concessionária — o SewerGEMS, da Bentley Systems.")

    h2(doc, "9.2. Horizonte de Projeto")
    add_para(doc,
             "As vazões de dimensionamento foram extraídas do Estudo de População e Vazão, que "
             "adotou horizonte de projeto de {{HORIZONTE_ANOS}} anos, de {{ANO_INI}} a {{ANO_FIM}}. "
             "Tanto os critérios hidráulicos quanto o dimensionamento das redes foram referidos à "
             "vazão máxima horária correspondente ao final do plano, de modo a assegurar capacidade "
             "e desempenho adequados durante todo o período de projeto.")

    h2(doc, "9.3. Vazões Específicas (Taxa de Contribuição Linear)")
    add_para(doc,
             "Para distribuir as contribuições ao longo da rede, recorreu-se ao conceito de vazão "
             "sanitária específica, ou taxa de contribuição linear (L/s·m). Tomou-se a vazão "
             "sanitária máxima horária do ano final do horizonte e rateou-se esse valor pela "
             "extensão total da rede coletora prevista, o que permite caracterizar a contribuição de "
             "forma contínua ao longo do sistema e embasar o dimensionamento das tubulações. A vazão "
             "assim obtida já incorpora a parcela de infiltração, segundo a taxa definida no item "
             "6.2.")
    add_table_caption(doc, "Estudo de vazões da {{SUBBACIA}} – síntese por ano (preencher do estudo)")
    make_table(
        doc,
        ["Ano", "Pop. Atend. (hab)", "Ext. Rede (m)", "Q infiltr. (L/s)",
         "Q média (L/s)", "Q máx. hor. (L/s)"],
        [
            ["{{ANO_INI}}", "{{POP_INI}}", "{{EXT_INI}}", "{{QINF_INI}}", "{{QMED_INI}}", "{{QMAX_INI}}"],
            ["...", "...", "...", "...", "...", "..."],
            ["{{ANO_FIM}}", "{{POP_FIM}}", "{{EXT_FIM}}", "{{QINF_FIM}}", "{{QMED_FIM}}", "{{QMAX_FIM}}"],
        ],
        col_widths=[2.0, 3.2, 2.8, 2.8, 2.8, 2.8],
    )
    add_table_source(doc, "Fonte: Estudo de População e Vazão – 2S Engenharia, 2026.")

    h2(doc, "9.4. Vazões Concentradas")
    add_para(doc,
             "Contribuições concentradas relevantes — provenientes de indústrias, estabelecimentos "
             "de ensino, unidades de saúde e demais grandes consumidores — só podem ser lançadas com "
             "segurança quando se conhecem em detalhe as particularidades de cada fonte: histórico de "
             "consumo, número de funcionários, alunos ou pacientes, períodos de funcionamento e "
             "localização. Na {{SUBBACIA}}, essas vazões foram tratadas a partir do levantamento de "
             "grandes consumidores e do acompanhamento de viabilidade de projetos hidrossanitários "
             "(PHS) repassado pela concessionária. {{OBS_VAZOES_CONCENTRADAS}}")

    # 9.5 - so quando o usuario aponta soleiras e/ou interferencias. Sem nenhum
    # dos dois, a secao inteira e omitida e o SewerGEMS sobe para 9.5.
    if TEM_SOLEIRAS and TEM_INTERF:
        _t95 = "9.5. Atendimento de Soleiras Negativas e Análise de Interferências"
    elif TEM_SOLEIRAS:
        _t95 = "9.5. Atendimento de Soleiras Negativas"
    elif TEM_INTERF:
        _t95 = "9.5. Análise de Interferências com Redes Existentes"
    else:
        _t95 = None

    if _t95:
        h2(doc, _t95)
    if TEM_SOLEIRAS:
        add_para(doc,
                 "Diz-se que uma soleira é negativa quando a cota de fundo do imóvel fica abaixo da cota "
                 "da rede coletora pública, situação em que o esgoto não escoa por gravidade pela ligação "
                 "convencional. Para ampliar ao máximo o atendimento, e a partir dos dados do "
                 "levantamento topográfico, empregou-se o critério de identificação de soleiras proposto "
                 "pela 2S Engenharia e Geotecnologia.")
        add_para(doc,
                 "Esse critério parte da cota de edificação (altitude do alicerce do imóvel). Desse "
                 "valor subtrai-se uma parcela fixa de 0,65 m, equivalente à profundidade típica da "
                 "ligação predial, obtendo-se a cota de fundo da ligação. Os imóveis foram inseridos no "
                 "modelo hidráulico como Property Connections, adotando-se a cota de edificação como cota "
                 "de terreno e a cota de fundo do imóvel como cota de fundo. Confrontando esses valores "
                 "com a profundidade máxima admitida para as redes, apurou-se o índice de atendimento da "
                 "bacia. A figura a seguir mostra a distribuição espacial das soleiras na {{SUBBACIA}}.")
        add_figure(doc, MAPA3, "Distribuição das soleiras na {{SUBBACIA}} (atendidas e não atendidas)")
        add_table_caption(doc, "Imóveis passíveis de atendimento na {{SUBBACIA}}")
        make_table(
            doc,
            ["Soleiras", "Número de Imóveis", "Percentual de Atendimento"],
            [
                ["Atendidas", "{{IMOVEIS_ATEND}}", "{{PCT_ATEND}}"],
                ["Não atendidas", "{{IMOVEIS_NAO_ATEND}}", "{{PCT_NAO_ATEND}}"],
                ["TOTAL DE IMÓVEIS", "{{IMOVEIS_TOTAL}}", "100,00%"],
            ],
            col_widths=[6.0, 4.5, 5.5],
            first_col_left=True,
        )
        add_table_source(doc)
    if TEM_INTERF:
        add_para(doc,
                 "Verificaram-se as possíveis interferências com outras infraestruturas. "
                 "Tomando o cadastro de redes existentes (drenagem, água e demais serviços) levantado "
                 "pela equipe de topografia, sobrepôs-se a ele o traçado da rede coletora projetada e "
                 "localizaram-se os pontos de cruzamento. Em cada um deles, examinaram-se as cotas e "
                 "profundidades relativas, considerando os diâmetros a partir da geratriz inferior dos "
                 "tubos, de modo a confirmar a compatibilidade das soluções adotadas.", space_before=6)
        add_para(doc,
                 "Da análise resultaram interferências com redes existentes de água ({{INTERF_AGUA}} "
                 "trechos) e de drenagem ({{INTERF_DREN}} trechos) cruzando o traçado projetado, todas "
                 "exigindo verificação de cotas e a definição das travessias correspondentes. A figura a "
                 "seguir localiza essas interferências na {{SUBBACIA}}.")
        add_figure(doc, MAPA5,
                   "Interferências identificadas (redes de água e drenagem existentes)")

    _n_sewer = "9.5" if not _t95 else "9.6"
    h2(doc, _n_sewer + ". Dados de Entrada no Software SewerGEMS")
    add_para(doc,
             "A modelagem hidráulica e o dimensionamento da rede coletora da {{SUBBACIA}} no "
             "SewerGEMS partiram do conjunto de parâmetros de entrada relacionados a seguir:",
             space_after=4)
    add_table_caption(doc, "Parâmetros de entrada para o dimensionamento da RCE")
    make_table(
        doc,
        ["Parâmetro", "Valor", "Unidade"],
        [
            ["Diâmetro mínimo (Dmin)", "{{DN_MIN}}", "mm"],
            ["Material", "{{MATERIAL_REDE}}", "–"],
            ["Coeficiente de Manning", "{{MANNING}}", "–"],
            ["Recobrimento mínimo", "{{RECOB_MIN}}", "m"],
            ["Declividade mínima", "{{DECL_MIN_PARAM}}", "m/m"],
            ["Tensão trativa mínima", "1,00", "Pa"],
            ["Lâmina d'água máxima", "75", "%"],
            ["Taxa de contribuição linear", "{{TAXA_LINEAR}}", "L/s·m"],
        ],
        col_widths=[8.0, 4.0, 4.0],
        first_col_left=True,
    )
    add_table_source(doc)

    # ====================================================================
    # 9. RESULTADOS - RESUMO DA REDE PROJETADA
    # ====================================================================
    new_section(doc)
    h1(doc, "10. RESULTADOS – RESUMO DA REDE PROJETADA")
    add_para(doc,
             "O detalhamento dos resultados trecho a trecho está reunido no Memorial de Cálculo "
             "(Planilha de Cálculo MC – documento {{CODIGO_MC}}). Neste capítulo apresenta-se apenas "
             "a consolidação dos quantitativos de tubulações e de dispositivos de inspeção da rede "
             "projetada para a {{SUBBACIA}}.")

    h2(doc, "10.1. Quantitativo de Rede por Diâmetro")
    add_table_caption(doc, "Quantitativo de rede da {{SUBBACIA}} por diâmetro")
    # tabela montada SO com os DN realmente presentes na rede (sem linhas em
    # branco de DN inexistentes). Para Amapora: apenas DN 150 mm.
    _dn_rows = [[dn, mat, ext] for (dn, mat, ext) in DN_ROWS]
    _dn_rows.append(["EXTENSÃO TOTAL DE REDES", "–", "{{EXT_TOTAL_REDE}}"])
    make_table(
        doc,
        ["Diâmetro", "Material", "Extensão (m)"],
        _dn_rows,
        col_widths=[6.0, 4.0, 6.0],
        first_col_left=True,
    )
    add_table_source(doc)

    h2(doc, "10.2. Quantitativo de Dispositivos de Inspeção")
    add_table_caption(doc, "Quantitativo de dispositivos de inspeção, tubos de queda e degraus da {{SUBBACIA}}")
    make_table(
        doc,
        ["Dispositivo", "Unidades", "Soma das alturas (m)"],
        [
            ["Poço de Visita (PV)", "{{QTD_PV}}", "–"],
            ["Terminal de Limpeza (TL)", "{{QTD_TL}}", "–"],
            ["Tubo de Queda (T.Q.)", "{{QTD_TQ}}", "{{TQ_SOMA_ALT}}"],
            ["Degrau", "{{QTD_DEGRAU}}", "{{DEGRAU_SOMA_ALT}}"],
            ["Terminal de Inspeção e Limpeza (TIL) {{QTD_TIL_NOTA}}", "{{QTD_TIL}}", "–"],
            ["Soleiras negativas (total)", "{{QTD_SOLEIRAS}}", "–"],
            ["TOTAL DE DISPOSITIVOS", "{{QTD_DISP_TOTAL}}", "–"],
        ],
        col_widths=[8.0, 3.5, 3.5],
        first_col_left=True,
    )
    add_table_source(doc)
    add_para(doc,
             "Foram previstos {{QTD_TQ}} tubos de queda (T.Q.), totalizando "
             "{{TQ_SOMA_ALT}} m de altura acumulada, e {{QTD_DEGRAU}} degraus, somando "
             "{{DEGRAU_SOMA_ALT}} m. Os tubos de queda atendem aos pontos em que o desnível "
             "entre a tubulação de montante e o fundo do poço de visita supera 0,50 m; os "
             "degraus, aos desníveis inferiores a esse limite.", space_before=4)

    h2(doc, "10.3. Distribuição por Faixa de Profundidade")
    add_para(doc,
             "A tabela a seguir apresenta a distribuição das estruturas da rede coletora por faixa "
             "de profundidade. A extensão de rede correspondente a cada faixa será consolidada a "
             "partir da Planilha de Cálculo (Memorial de Cálculo – MC).", space_after=4)
    add_table_caption(doc, "Distribuição de estruturas da rede por faixa de profundidade")
    make_table(
        doc,
        ["Faixa de Profundidade", "Nº de Estruturas"],
        [
            ["Até 1,25 m", "{{EXT_PROF_125}}"],
            ["De 1,25 a 2,00 m", "{{EXT_PROF_200}}"],
            ["De 2,00 a 3,00 m", "{{EXT_PROF_300}}"],
            ["De 3,00 a 4,00 m", "{{EXT_PROF_400}}"],
            ["Acima de 4,00 m", "{{EXT_PROF_500}}"],
        ],
        col_widths=[9.0, 6.0],
        first_col_left=True,
    )
    add_table_source(doc)

    h2(doc, "10.4. Método Executivo")
    add_para(doc,
             "A construção das redes foi prevista com os métodos mais adequados a cada situação de "
             "pavimentação, profundidade e interferência: a Vala a Céu Aberto (VCA) e, quando "
             "justificável, o Método Não Destrutivo (MND). A tabela a seguir resume a extensão "
             "atribuída a cada método executivo.", space_after=4)
    add_table_caption(doc, "Extensão de rede por método executivo")
    make_table(
        doc,
        ["Método Executivo", "Extensão (m)", "Percentual (%)"],
        [
            ["Vala a Céu Aberto (VCA)", "{{EXT_VCA}}", "{{PCT_VCA}}"],
            ["Método Não Destrutivo (MND)", "{{EXT_MND}}", "{{PCT_MND}}"],
            ["EXTENSÃO TOTAL DE REDES", "{{EXT_TOTAL}}", "100,00%"],
        ],
        col_widths=[7.0, 4.0, 4.0],
        first_col_left=True,
    )
    add_table_source(doc)

    # ====================================================================
    # 10. CONSIDERACOES FINAIS
    # ====================================================================
    new_section(doc)
    h1(doc, "11. CONSIDERAÇÕES FINAIS")
    add_para(doc,
             "Ao longo deste Memorial Descritivo foram expostos a concepção, os critérios e os "
             "parâmetros que nortearam o dimensionamento da rede coletora de esgoto da {{SUBBACIA}}, "
             "no município de {{MUNICIPIO}} – {{UF}}, sempre em conformidade com as normas da ABNT, "
             "com os manuais e diretrizes da SANEPAR e com as premissas acertadas com a ACCIONA.")
    add_para(doc,
             "O cálculo hidráulico, processado no SewerGEMS para a vazão máxima horária ao final do "
             "horizonte de projeto, respeitou os limites de tensão trativa mínima, velocidade, "
             "lâmina d'água e recobrimento. As soluções resultantes garantem o bom funcionamento "
             "hidráulico e a exequibilidade construtiva do sistema" +
             (", alcançando índice de atendimento da bacia de {{PCT_ATEND}}." if TEM_SOLEIRAS else "."))
    if TEM_SOLEIRAS:
        add_para(doc,
                 "As soleiras negativas que, por conta da limitação de profundidade das redes, não "
                 "puderam ser atendidas por gravidade foram devidamente identificadas e justificadas; "
                 "seu atendimento deverá valer-se de soluções individuais específicas, de acordo com as "
                 "orientações da concessionária.")
    add_para(doc,
             "Por fim, recomenda-se que a execução das obras siga estritamente os critérios "
             "construtivos aqui descritos, as especificações do MOS/SANEPAR e os detalhes das "
             "plantas e perfis do projeto executivo, de modo a assegurar a qualidade, a durabilidade "
             "e o desempenho do sistema durante toda a sua vida útil.")

    # ====================================================================
    # 11. ANEXOS
    # ====================================================================
    new_section(doc)
    h1(doc, "12. ANEXOS")
    add_para(doc, "Acompanham este Memorial Descritivo, como peças complementares, os seguintes documentos:", space_after=4)
    add_bullet(doc, "Anexo I – Planilha de Cálculo (Memorial de Cálculo – MC), documento {{CODIGO_MC}};")
    add_bullet(doc, "Anexo II – Plantas e Perfis da Rede Coletora (pranchas do projeto executivo);")
    add_bullet(doc, "Anexo III – ART – Anotação de Responsabilidade Técnica nº {{ART}};")
    add_bullet(doc, "Anexo IV – Relatório Técnico de Levantamento Topográfico;")
    add_bullet(doc, "Anexo V – Boletins de Sondagem / Investigação Geotécnica;")
    add_bullet(doc, "Anexo VI – Estudo de População e Vazão.")

    # --- preenche as LISTAS de figuras e tabelas LITERALMENTE (sem F9) ---
    # insere de tras para frente para manter a ordem crescente apos a ancora
    for num, leg in reversed(FIGURAS):
        _insert_list_paragraph_after(doc, anchor_fig, "Figura", num, leg)
    for num, leg in reversed(TABELAS):
        _insert_list_paragraph_after(doc, anchor_tab, "Tabela", num, leg)

    # substitui os {{PLACEHOLDERS}} pelos dados reais de Amapora; campos
    # ausentes (vazao/populacao/coeficientes/fluxograma) saem em VERMELHO.
    apply_substitutions(doc)

    # atualiza os campos (SUMARIO/PAGE) automaticamente ao abrir no Word,
    # sem necessidade de F9.
    _set_update_fields_on_open(doc)

    # salvar
    doc.save(SAIDA)
    return doc


# ----------------------------------------------------------------------------
# Configuracao dirigida por JSON (integracao Nexus)
# ----------------------------------------------------------------------------
def _resolver_mapas(mapas_dir):
    """Resolve os 6 mapas a partir de um diretorio. Aceita os nomes padrao
    (Mapa1_Localizacao.png ... Mapa6_3D_Topografia.png) ou, na falta, o
    primeiro PNG que comece com 'Mapa<N>'."""
    global MAPA1, MAPA2, MAPA3, MAPA4, MAPA5, MAPA6, MAPA7
    if not mapas_dir or not os.path.isdir(mapas_dir):
        return
    nomes = {
        1: "Mapa1_Localizacao.png", 2: "Mapa2_SubBacia_SB02_Rede.png",
        3: "Mapa3_Soleiras_SB02.png", 4: "Mapa4_Calor_Topografia.png",
        5: "Mapa5_Interferencias.png", 6: "Mapa6_3D_Topografia.png",
        7: "Mapa7_Sondagem.png",
    }
    arquivos = os.listdir(mapas_dir)

    def achar(n):
        p = os.path.join(mapas_dir, nomes[n])
        if os.path.exists(p):
            return p
        pref = "mapa%d" % n
        for a in sorted(arquivos):
            if a.lower().startswith(pref) and a.lower().endswith(".png"):
                return os.path.join(mapas_dir, a)
        return p
    MAPA1, MAPA2, MAPA3 = achar(1), achar(2), achar(3)
    MAPA4, MAPA5, MAPA6 = achar(4), achar(5), achar(6)
    MAPA7 = achar(7)


def _gerar_fluxograma_do_config(cfg, out_png):
    """Gera o PNG do fluxograma a partir das etapas do config, importando
    gerar_fluxograma.gerar(). Retorna o caminho do PNG ou None se falhar."""
    flux = cfg.get("fluxograma") or {}
    etapas = flux.get("etapas") or []
    if not etapas:
        return None
    try:
        import gerar_fluxograma as gf
        proj = cfg.get("projeto") or {}
        fcfg = {
            "municipio": proj.get("municipio", "Amaporã"),
            "subbacia": proj.get("subbacia", ""),
            "etapas": etapas,
        }
        if flux.get("titulo"):
            fcfg["titulo"] = flux["titulo"]
        gf.gerar(fcfg, out_png)
        return out_png if os.path.exists(out_png) else None
    except Exception as e:
        sys.stderr.write("Falha ao gerar fluxograma: %s\n" % e)
        return None


def _tem_brutos(cfg):
    """True se o config aponta arquivos brutos (OSE/TXT/soleiras/interferencias)."""
    return any(cfg.get(k) for k in ("ose", "ose_xlsx", "txt_dir", "soleiras",
                                    "interferencias", "modelo", "modelo_sqlite"))


def _gerar_mapas(geo_dir, mapas_dir, cfg):
    """Gera os 6 PNGs dos mapas rodando mapa1..6 como subprocessos, com os
    caminhos passados via variaveis de ambiente (MEMORIAL_*). Roda na pasta do
    script (cwd) p/ os imports de cartolib resolverem. Retorna mapas_dir."""
    _ensure(mapas_dir)
    cache_dir = _ensure(os.path.join(os.path.dirname(mapas_dir), "cache"))
    # semeia o cache de contexto (malhas municipais IBGE p/ o mapa1) a partir de
    # um cache existente: SCRIPT_DIR/cache (bundle) ou BASE/cache (legado).
    for legado_cache in (os.path.join(SCRIPT_DIR, "cache"), os.path.join(BASE, "cache")):
        if os.path.isdir(legado_cache):
            for fn in os.listdir(legado_cache):
                if fn.endswith((".geojson", ".json")):
                    dst = os.path.join(cache_dir, fn)
                    if not os.path.exists(dst):
                        try:
                            shutil.copyfile(os.path.join(legado_cache, fn), dst)
                        except Exception:
                            pass

    proj = cfg.get("projeto") or {}
    env = dict(os.environ)
    env["MEMORIAL_GEO"] = geo_dir
    env["MEMORIAL_MAPAS"] = mapas_dir
    env["MEMORIAL_CACHE"] = cache_dir
    if cfg.get("txt_dir"):
        env["MEMORIAL_TXT_DIR"] = cfg["txt_dir"]
    if cfg.get("interferencias"):
        env["MEMORIAL_INTERF"] = cfg["interferencias"]
    # diretorio com os shapes de interferencia ja classificados (nomes canonicos),
    # gerado pela extracao em geo/interferencias_shp — tem prioridade no mapa5.
    _interf_shp = os.path.join(geo_dir, "interferencias_shp")
    if os.path.isdir(_interf_shp):
        env["MEMORIAL_INTERF_SHP"] = _interf_shp
    ibge = cfg.get("codigo_ibge") or proj.get("codigo_ibge")
    if ibge:
        env["MEMORIAL_IBGE"] = str(ibge)
    env["MEMORIAL_MUNICIPIO"] = proj.get("municipio") or "Amaporã"
    if cfg.get("sondagem_xlsx"):
        env["MEMORIAL_SONDAGEM_XLSX"] = cfg["sondagem_xlsx"]
    env["PYTHONIOENCODING"] = "utf-8"

    import subprocess
    # cada item: (script, png esperado). Usado para conferir o que saiu de fato.
    scripts = [
        ("mapa1_localizacao.py",  "Mapa1_Localizacao.png"),
        ("mapa2_subbacia.py",     "Mapa2_SubBacia_SB02_Rede.png"),
        ("mapa3_soleiras.py",     "Mapa3_Soleiras_SB02.png"),
        ("mapa4_calor.py",        "Mapa4_Calor_Topografia.png"),
        ("mapa5_interferencias.py", "Mapa5_Interferencias.png"),
        ("mapa6_3d.py",           "Mapa6_3D_Topografia.png"),
        ("mapa7_sondagem.py",     "Mapa7_Sondagem.png"),
    ]
    gerados = []   # PNGs que existem ao final
    falhas = []    # (script, motivo)
    for sc, png in scripts:
        spath = os.path.join(SCRIPT_DIR, sc)
        out_png = os.path.join(mapas_dir, png)
        if not os.path.exists(spath):
            sys.stderr.write("[mapas] Script nao encontrado: %s\n" % spath)
            falhas.append((sc, "script ausente"))
            continue
        # ROBUSTEZ: cada mapa roda isolado; se um falhar (ex.: mapa5 sem
        # interferencias), capturamos, LOGAMOS e CONTINUAMOS os demais — o
        # pipeline nao aborta por causa de um mapa.
        try:
            r = subprocess.run([sys.executable, spath], cwd=SCRIPT_DIR, env=env,
                               capture_output=True, text=True, encoding="utf-8")
            if r.returncode != 0:
                sys.stderr.write(
                    "[mapas] FALHOU %s (rc=%s) — continuando os demais:\n%s\n%s\n"
                    % (sc, r.returncode, (r.stdout or "").strip(),
                       (r.stderr or "").strip()))
                falhas.append((sc, "returncode=%s" % r.returncode))
            else:
                # propaga o stdout do mapa (avisos como "interferencias nao
                # encontradas -> mapa5 sem camadas/omitido") para o log.
                if r.stdout and r.stdout.strip():
                    sys.stderr.write("[mapas] %s: %s\n" % (sc, r.stdout.strip()))
        except Exception as e:
            sys.stderr.write("[mapas] ERRO ao rodar %s: %s — continuando.\n" % (sc, e))
            falhas.append((sc, str(e)))
        # confere o resultado pelo PNG (independente do returncode): pode ter
        # sido OMITIDO de proposito (mapa5 sem interferencias) ou caido fora.
        if os.path.exists(out_png):
            gerados.append(png)
        else:
            sys.stderr.write("[mapas] %s nao produziu %s (omitido/sem dados).\n"
                             % (sc, png))
    sys.stderr.write("[mapas] gerados %d/%d: %s%s\n" % (
        len(gerados), len(scripts), ", ".join(gerados) or "(nenhum)",
        ("; falhas: " + ", ".join("%s(%s)" % f for f in falhas)) if falhas else ""))
    return mapas_dir


def _ensure(p):
    os.makedirs(p, exist_ok=True)
    return p


def _executar_extracao(cfg):
    """Roda a extracao parametrica (extrair_dados.extrair) a partir dos brutos do
    config e gera os 6 mapas. Atualiza cfg com os caminhos resultantes
    (dados_json, dados_extra_json, rede_geojson, mapas_dir). Retorna o resumo."""
    import extrair_dados
    # out_dir: ao lado da saida (ou tempdir)
    out_dir = (cfg.get("out_dir") or
               (os.path.join(os.path.dirname(os.path.abspath(cfg["saida"])), "_extracao")
                if cfg.get("saida") else os.path.join(tempfile.gettempdir(), "memorial_extracao")))
    _ensure(out_dir)
    cfg["out_dir"] = out_dir
    resumo = extrair_dados.extrair(cfg)
    # liga os artefatos gerados de volta no config
    cfg["dados_json"] = resumo["dados_json"]
    cfg["dados_extra_json"] = resumo["dados_extra_json"]
    cfg["rede_geojson"] = resumo["rede_geojson"]
    # gera mapas (a menos que mapas_dir ja exista com PNGs e nao se peca regerar)
    mapas_dir = cfg.get("mapas_dir") or os.path.join(out_dir, "mapas")
    pngs = ([p for p in os.listdir(mapas_dir)
             if p.lower().endswith(".png")] if os.path.isdir(mapas_dir) else [])
    if cfg.get("gerar_mapas", True) and len(pngs) < 6:
        _gerar_mapas(resumo["geo_dir"], mapas_dir, cfg)
    cfg["mapas_dir"] = mapas_dir
    return resumo


def apply_config(cfg):
    """Aplica o JSON de configuracao aos globais de caminho/dados deste modulo.
    Estrutura esperada (todos os campos opcionais; faltantes caem no Amapora):

    {
      "saida": "C:/.../Memorial.docx",      # caminho do .docx de saida
      "base": "C:/.../projeto",             # pasta com dados/mapas (modo legado)
      "dados_json": "...",                  # JSON de quantitativos (dados_amapora.json)
      "dados_extra_json": "...",            # JSON extra (TQ/DN/GNSS)
      "rede_geojson": "...",                # rede_trechos.geojson (footprint bacia)
      "mapas_dir": "C:/.../mapas",          # pasta com os 6 mapas PNG
      "template": "...",                    # template_2s.docx (default = bundle)
      "projeto": { municipio, uf, subbacia, codigo_doc, ... },
      "art":     { eng_resp, crea, art, ... },
      "fluxograma": { "etapas": [ {tipo, descricao}, ... ], "titulo": "" }
    }
    """
    global CFG, BASE, MAPAS, FLUXOGRAMA, FOTOS, DADOS_JSON, DADOS_EXTRA_JSON
    global SAIDA, TEMPLATE, REDE_GEOJSON, WORKDIR, TEM_SOLEIRAS, TEM_INTERF
    CFG = cfg or {}

    # Modo legado (sem --config / sem brutos): mantem a secao 9.5 como antes.
    # Com config: a secao so sai se o usuario apontar soleiras / interferencias.
    if _tem_brutos(cfg):
        TEM_SOLEIRAS = bool(cfg.get("soleiras"))
        TEM_INTERF = bool(cfg.get("interferencias"))

    if cfg.get("base"):
        BASE = cfg["base"]
        MAPAS = os.path.join(BASE, "mapas")
        FOTOS = os.path.join(BASE, "fotos_campo")

    # diretorio de trabalho gravavel (ao lado da saida, senao tempdir)
    if cfg.get("saida"):
        SAIDA = cfg["saida"]
        out_dir = os.path.dirname(os.path.abspath(SAIDA))
        if out_dir and os.path.isdir(out_dir):
            WORKDIR = out_dir

    if cfg.get("template") and os.path.exists(cfg["template"]):
        TEMPLATE = cfg["template"]

    # ---- ETAPA DE EXTRACAO (brutos -> dados.json + mapas) -------------------
    # Se o config trouxer arquivos brutos (OSE/TXT/soleiras/interferencias) e
    # ainda nao houver dados_json pronto (ou se 'extrair': true), roda o pipeline
    # de extracao parametrico ANTES do build, fechando o fluxo "brutos -> .docx".
    _dados_existe = bool(cfg.get("dados_json") and os.path.exists(cfg["dados_json"]))
    _pedir_extrair = bool(cfg.get("extrair"))
    if _tem_brutos(cfg) and (_pedir_extrair or not _dados_existe):
        try:
            res = _executar_extracao(cfg)
            sys.stderr.write("[extracao] OK: %d OSE, %s m, %d PV, %d TL, TQ %d, deg %d\n" % (
                res.get("n_oses", 0), res.get("extensao_total_m", "?"),
                res.get("n_PV", 0), res.get("n_TL", 0),
                res.get("n_TQ", 0), res.get("n_DEGRAU", 0)))
        except Exception as e:
            import traceback
            sys.stderr.write("[extracao] FALHA:\n%s\n" % traceback.format_exc())
            raise RuntimeError("Falha na extracao dos dados brutos: %s" % e)

    # dados (quantitativos / extra / rede)
    if cfg.get("dados_json"):
        DADOS_JSON = cfg["dados_json"]
    if cfg.get("dados_extra_json"):
        DADOS_EXTRA_JSON = cfg["dados_extra_json"]
    if cfg.get("rede_geojson"):
        REDE_GEOJSON = cfg["rede_geojson"]
    elif cfg.get("base"):
        REDE_GEOJSON = os.path.join(BASE, "geo_amapora", "rede_trechos.geojson")

    # mapas
    _resolver_mapas(cfg.get("mapas_dir") or MAPAS)

    # fluxograma: gera do config (etapas) num PNG temporario; senao usa o do BASE
    flux_png = os.path.join(WORKDIR, "_fluxograma_memorial.png")
    gerado = _gerar_fluxograma_do_config(cfg, flux_png)
    if gerado:
        FLUXOGRAMA = gerado
    elif cfg.get("fluxograma_png") and os.path.exists(cfg["fluxograma_png"]):
        FLUXOGRAMA = cfg["fluxograma_png"]
    # senao mantem o default (BASE/Fluxograma_Sistema.png)


def main():
    ap = argparse.ArgumentParser(
        description="Gerador de Memorial Descritivo RCE (2S) — dirigido por JSON")
    ap.add_argument("--config", help="JSON com caminhos + dados do projeto")
    ap.add_argument("--saida", help="Caminho do .docx de saida (sobrescreve config)")
    args = ap.parse_args()

    if args.config:
        try:
            with open(args.config, "r", encoding="utf-8") as f:
                cfg = json.load(f)
        except Exception as e:
            print(json.dumps({"ok": False, "erro": "Falha ao ler config: %s" % e}))
            return 1
        if args.saida:
            cfg["saida"] = args.saida
        try:
            apply_config(cfg)
        except Exception as e:
            print(json.dumps({"ok": False, "erro": "Falha ao aplicar config: %s" % e}))
            return 1

        # Os dados quantitativos (dados_json) sao obrigatorios. Quando o config
        # traz arquivos brutos, apply_config() ja rodou a EXTRACAO parametrica e
        # gerou o dados_json + os 6 mapas. Se mesmo assim nao houver dados, erra.
        if not os.path.exists(DADOS_JSON):
            print(json.dumps({"ok": False, "erro":
                "Arquivo de dados (dados_json) nao encontrado: %s. "
                "Informe os arquivos brutos (ose/txt_dir/soleiras/interferencias) "
                "para extracao automatica, ou um 'dados_json' ja pronto." % DADOS_JSON}))
            return 1

    try:
        build()
    except Exception as e:
        import traceback
        sys.stderr.write(traceback.format_exc())
        print(json.dumps({"ok": False, "erro": str(e)}))
        return 1

    # saida JSON (ultima linha do stdout) — padrao consumido pelo Nexus
    print(json.dumps({
        "ok": True,
        "saida": SAIDA,
        "figuras": _fig_n[0],
        "tabelas": _tab_n[0],
    }))
    return 0


if __name__ == "__main__":
    sys.exit(main())
